"""
Build the KU-branded SPMD Minor → Fall 2027 ADEK Roadmap Word document.

Reads  ../_data.json
Writes docs/SPMD_Fall2027_Roadmap.docx

Audience: ADEK reviewers
QF Emirates Level 6 framing
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths & constants
# ---------------------------------------------------------------------------
HERE = Path(__file__).resolve().parent
DATA_PATH = HERE.parent / "_data.json"
OUT_PATH = HERE / "SPMD_Fall2027_Roadmap.docx"

GEN_DATE = "2026-05-18"
FOOTER_TEXT = "BSCND Sports Medicine Roadmap · Fall 2027 · Generated 2026-05-18 · ADEK-aligned"

# KU brand
KU_NAVY = RGBColor(0x00, 0x3D, 0x7A)
KU_NAVY_HEX = "003D7A"
KU_GOLD = RGBColor(0xC8, 0xA0, 0x4A)
KU_GOLD_HEX = "C8A04A"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_GREY = RGBColor(0x1F, 0x29, 0x37)
BODY_GREY_HEX = "1F2937"
AMBER = RGBColor(0xD9, 0x77, 0x06)
AMBER_HEX = "D97706"
AMBER_LIGHT_HEX = "FFF7ED"
GREEN = RGBColor(0x1B, 0x8A, 0x4E)
ALT_ROW_HEX = "F3F4F6"

FONT_NAME = "Calibri"
H1_PT = 20
H2_PT = 14
H3_PT = 11
BODY_PT = 10
SMALL_PT = 9
TINY_PT = 8


# ---------------------------------------------------------------------------
# OXML helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _table_default_borders(table, color: str = "BFBFBF", sz: int = 4) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tbl_borders.append(b)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_borders)


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_horizontal_rule(paragraph, color_hex: str, sz: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def _add_page_break(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _set_font(run, size_pt=None, color=None, bold=None, italic=None) -> None:
    run.font.name = FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    rfonts = r_pr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        r_pr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def _para(container, text: str, *, size_pt=BODY_PT, color=BODY_GREY,
          bold=False, italic=False, align=None,
          space_after_pt: float = 4, space_before_pt: float = 0,
          line_spacing: float | None = None):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.space_before = Pt(space_before_pt)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    _set_font(run, size_pt=size_pt, color=color, bold=bold, italic=italic)
    return p


def _set_col_widths(table, widths_inches: list) -> None:
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    existing = tbl_pr.find(qn("w:tblLayout"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(layout)
    for row in table.rows:
        for idx, w in enumerate(widths_inches):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(w)


def _cell_text(cell, text: str, *, size_pt=BODY_PT, color=BODY_GREY,
               bold=False, italic=False, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    _set_font(run, size_pt=size_pt, color=color, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _header_row(table, headers: list, fill_hex: str = KU_NAVY_HEX) -> None:
    row = table.rows[0]
    _set_repeat_header(row)
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _shade_cell(cell, fill_hex)
        _cell_text(cell, h, size_pt=BODY_PT, color=WHITE, bold=True)


def _alternate_row_shade(table, start_row: int = 1) -> None:
    for i, row in enumerate(table.rows):
        if i < start_row:
            continue
        if (i - start_row) % 2 == 1:
            for cell in row.cells:
                _shade_cell(cell, ALT_ROW_HEX)


def _add_gold_left_border(table) -> None:
    for row in table.rows:
        cell = row.cells[0]
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge_name in ("left",):
            existing = tc_borders.find(qn(f"w:{edge_name}"))
            if existing is not None:
                tc_borders.remove(existing)
            b = OxmlElement(f"w:{edge_name}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "18")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), KU_GOLD_HEX)
            tc_borders.append(b)


def _add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    _set_font(run, size_pt=SMALL_PT, color=BODY_GREY)
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r.append(fld_sep)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_end)


# ---------------------------------------------------------------------------
# Section / page setup
# ---------------------------------------------------------------------------

def _setup_section(section) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.different_first_page_header_footer = True


def _add_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p._p.getparent().remove(p._p)
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("KHALIFA UNIVERSITY  ·  College of Medicine and Health Sciences  ·  ADEK Submission")
    _set_font(run, size_pt=SMALL_PT, color=KU_NAVY, bold=True)
    _add_horizontal_rule(p, KU_GOLD_HEX, sz=8)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._p.getparent().remove(p._p)
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=2)
    left_run = fp.add_run(FOOTER_TEXT)
    _set_font(left_run, size_pt=SMALL_PT, color=BODY_GREY)
    fp.add_run("\t")
    run_page = fp.add_run("Page ")
    _set_font(run_page, size_pt=SMALL_PT, color=BODY_GREY)
    _add_field(fp, "PAGE")
    run_of = fp.add_run(" of ")
    _set_font(run_of, size_pt=SMALL_PT, color=BODY_GREY)
    _add_field(fp, "NUMPAGES")

    # Blank first-page header/footer
    fph = section.first_page_header
    fph.is_linked_to_previous = False
    for p in list(fph.paragraphs):
        p._p.getparent().remove(p._p)
    fph.add_paragraph()
    fpf = section.first_page_footer
    fpf.is_linked_to_previous = False
    for p in list(fpf.paragraphs):
        p._p.getparent().remove(p._p)
    fpf.add_paragraph()


# ---------------------------------------------------------------------------
# Heading helpers
# ---------------------------------------------------------------------------

def _h1(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_font(run, size_pt=H1_PT, color=KU_NAVY, bold=True)
    _add_horizontal_rule(p, KU_GOLD_HEX, sz=12)


def _h2(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, size_pt=H2_PT, color=KU_NAVY, bold=True)
    _add_horizontal_rule(p, KU_GOLD_HEX, sz=6)


def _h3(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_font(run, size_pt=H3_PT, color=KU_NAVY, bold=True)


def _bullet(doc, text: str, *, size_pt=BODY_PT, color=BODY_GREY, italic=False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_font(run, size_pt=size_pt, color=color, italic=italic)


def _numbered(doc, text: str, *, size_pt=BODY_PT, color=BODY_GREY) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_font(run, size_pt=size_pt, color=color)


# ---------------------------------------------------------------------------
# 1. Cover
# ---------------------------------------------------------------------------

def build_cover(doc, data: dict) -> None:
    program = data["program"]

    # Navy header band
    band = doc.add_table(rows=2, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    _set_col_widths(band, [6.5])
    _table_default_borders(band, color="FFFFFF", sz=0)
    top_cell = band.rows[0].cells[0]
    bot_cell = band.rows[1].cells[0]
    _shade_cell(top_cell, KU_NAVY_HEX)
    _shade_cell(bot_cell, KU_GOLD_HEX)

    top_cell.text = ""
    p = top_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("KHALIFA UNIVERSITY OF SCIENCE AND TECHNOLOGY")
    _set_font(run, size_pt=16, color=WHITE, bold=True)

    bot_cell.text = ""
    bp = bot_cell.paragraphs[0]
    bp.paragraph_format.space_before = Pt(1)
    bp.paragraph_format.space_after = Pt(1)
    rr = bp.add_run(" ")
    _set_font(rr, size_pt=2, color=KU_GOLD)

    for _ in range(5):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("College of Medicine and Health Sciences")
    _set_font(run, size_pt=14, color=BODY_GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("BSc Clinical Nutrition & Dietetics (BSCND)")
    _set_font(run, size_pt=12, color=KU_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("SPMD Minor → Fall 2027 ADEK Roadmap")
    _set_font(run, size_pt=26, color=KU_NAVY, bold=True)

    # Gold rule
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_horizontal_rule(sep, KU_GOLD_HEX, sz=12)
    sep.paragraph_format.space_after = Pt(12)

    # Sub-title line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("QF Emirates Level 6 Programme Authorization Submission")
    _set_font(run, size_pt=12, color=BODY_GREY, italic=True)

    for _ in range(5):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Metadata table
    meta_rows = [
        ("Programme", program["full_name"]),
        ("Institution", program["institution"]),
        ("College", program["college"]),
        ("Target Launch", program["launch"]),
        ("Date", GEN_DATE),
        ("Prepared by", "Dr. Carlo Raj"),
        ("Audience", "ADEK Reviewers"),
    ]
    meta = doc.add_table(rows=len(meta_rows), cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    _set_col_widths(meta, [2.0, 4.5])
    _table_default_borders(meta, color="FFFFFF", sz=0)
    for i, (k, v) in enumerate(meta_rows):
        _cell_text(meta.rows[i].cells[0], k, size_pt=SMALL_PT, color=KU_NAVY, bold=True)
        _cell_text(meta.rows[i].cells[1], v, size_pt=SMALL_PT, color=BODY_GREY)


# ---------------------------------------------------------------------------
# 2. Executive Summary
# ---------------------------------------------------------------------------

def build_executive_summary(doc, data: dict) -> None:
    _h1(doc, "Executive Summary")

    _para(doc,
          "This document presents the ADEK Programme Authorization roadmap for the proposed Sports "
          "Medicine Minor (SPMD 301–304) and Sports & Fitness Coaching Concentration "
          "(SPMD 305–308) within the Bachelor of Science in Clinical Nutrition and Dietetics "
          "(BSCND) at Khalifa University College of Medicine and Health Sciences (CMHS).",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    _h2(doc, "Two-Track Approach")
    _para(doc,
          "The roadmap adopts a dual-track strategy responsive to both institutional governance and "
          "UAE professional-registration requirements. Track 1 (SPMD 301–304: Sports Medicine "
          "Minor) consolidates four existing syllabi awaiting UGCC review. Track 2 (SPMD 305–308: "
          "Sports & Fitness Coaching Concentration) introduces four new courses proposed in response "
          "to Dean Dr. Habiba Alsafar’s May 2026 inquiry, closing all ICREPs Level 3 PT "
          "competency gaps and enabling REPs UAE Personal Trainer registration for graduates.",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    _h2(doc, "Fall 2027 Launch Window")
    _para(doc,
          "The five-quarter implementation roadmap (Q3 2026 – Q3 2027) targets first-class "
          "enrolment in SPMD 301 and SPMD 305 in September 2027. Key governance gates include: "
          "bundled UGCC submission (Q3 2026), ADEK Program Authorization for the minor (Q4 2026), "
          "MoHESR filing (Q1 2027), and KU Senate/Board approval of the concentration (Q2 2027).",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    _h2(doc, "Regulator Alignment")
    _para(doc,
          "The concentration is aligned to six UAE regulatory bodies: REPs UAE (primary professional "
          "registration), ADSC and DSC (emirate-level enforcement), GAS (federal strategic "
          "framework), DoH Abu Dhabi (parallel Dietitian licence pathway), and MoHRE (commercial "
          "employment compliance). Alignment letters from REPs UAE / ADSC are targeted for the "
          "Q3 2027 Go-Live gate.",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    _h2(doc, "Key Programme Facts")
    facts = doc.add_table(rows=5, cols=2)
    facts.autofit = False
    _set_col_widths(facts, [2.5, 4.0])
    _table_default_borders(facts, color="BFBFBF", sz=4)
    fact_rows = [
        ("Baseline credits (BSCND)", str(data["program"]["total_credits_baseline"])),
        ("Credits with Concentration", str(data["program"]["total_credits_with_concentration"])),
        ("Minor credits", "12 (SPMD 301–304)"),
        ("Concentration credits", "12 (SPMD 305–308)"),
        ("QF Emirates Level", str(data["program"]["qf_emirates_level"])),
    ]
    for i, (k, v) in enumerate(fact_rows):
        _shade_cell(facts.rows[i].cells[0], KU_NAVY_HEX)
        _cell_text(facts.rows[i].cells[0], k, size_pt=BODY_PT, color=WHITE, bold=True)
        _cell_text(facts.rows[i].cells[1], v, size_pt=BODY_PT, color=BODY_GREY)
    _add_gold_left_border(facts)


# ---------------------------------------------------------------------------
# 3. Two-Track Programme Structure
# ---------------------------------------------------------------------------

def build_track_structure(doc, data: dict) -> None:
    _h1(doc, "Two-Track Programme Structure")
    tracks = data["tracks"]

    for track in sorted(tracks, key=lambda t: t["id"]):
        label = "Track 1 — Sports Medicine Minor" if track["id"] == "minor" else "Track 2 — Sports & Fitness Coaching Concentration"
        fill = ALT_ROW_HEX if track["id"] == "minor" else AMBER_LIGHT_HEX
        _h2(doc, label)

        tbl = doc.add_table(rows=5, cols=2)
        tbl.autofit = False
        _set_col_widths(tbl, [2.0, 4.5])
        _table_default_borders(tbl, color="BFBFBF", sz=4)
        rows = [
            ("Full name", track["name"]),
            ("Courses", ", ".join(track["courses"])),
            ("Total credits", str(track["credits_total"])),
            ("Status", track["status"]),
            ("ICREPs eligible", "Yes — closes REPs UAE PT registration gaps" if track.get("icreps_eligible") else "No — didactic athlete-centred curriculum"),
        ]
        for i, (k, v) in enumerate(rows):
            _shade_cell(tbl.rows[i].cells[0], KU_NAVY_HEX)
            _cell_text(tbl.rows[i].cells[0], k, size_pt=BODY_PT, color=WHITE, bold=True)
            _cell_text(tbl.rows[i].cells[1], v, size_pt=BODY_PT, color=BODY_GREY)
        _add_gold_left_border(tbl)

        _para(doc, track["note"], size_pt=SMALL_PT, color=BODY_GREY, italic=True, space_after_pt=10)

    _h2(doc, "Side-by-Side Comparison")
    cmp_tbl = doc.add_table(rows=6, cols=3)
    cmp_tbl.autofit = False
    _set_col_widths(cmp_tbl, [2.0, 2.25, 2.25])
    _table_default_borders(cmp_tbl, color="BFBFBF", sz=4)
    _header_row(cmp_tbl, ["Attribute", "SPMD Minor (301–304)", "Concentration (305–308)"])
    cmp_rows = [
        ("Course codes", "SPMD 301, 302, 303, 304", "SPMD 305, 306, 307, 308"),
        ("Credits", "12", "12"),
        ("Target student", "Nutrition students wishing to deepen sports science knowledge; athlete-centred; no fitness-coach credential", "Nutrition students targeting REPs UAE Personal Trainer / Group Fitness Instructor registration"),
        ("REPs UAE eligible", "No", "Yes — ICREPs Level 3 PT pathway"),
        ("Governance status", "Pending UGCC review", "Proposed — syllabi not yet authored"),
    ]
    for i, row in enumerate(cmp_rows, start=1):
        for j, val in enumerate(row):
            _cell_text(cmp_tbl.rows[i].cells[j], val, size_pt=SMALL_PT, color=BODY_GREY)
    _alternate_row_shade(cmp_tbl)
    _add_gold_left_border(cmp_tbl)


# ---------------------------------------------------------------------------
# 4. Course Catalogue
# ---------------------------------------------------------------------------

def build_course_catalogue(doc, data: dict) -> None:
    _h1(doc, "Course Catalogue — SPMD 301–308")
    courses = data["courses"]
    plos = {p["id"]: p["label"] for p in data["bscnd_plos"]}

    for code in sorted(courses.keys()):
        c = courses[code]
        is_inferred = c.get("_inferred", False)
        track_label = "Concentration (SPMD 305–308)" if c["track"] == "concentration" else "Minor (SPMD 301–304)"
        heading = f"{code}: {c['title']}  [{track_label}]"
        _h2(doc, heading)

        # Proposed tag for inferred courses
        if is_inferred:
            warn_p = doc.add_paragraph()
            warn_p.paragraph_format.space_after = Pt(4)
            warn_run = warn_p.add_run("⚠  PROPOSED — pending faculty validation")
            _set_font(warn_run, size_pt=BODY_PT, color=AMBER, bold=True)

        # Course metadata table
        prereq_str = "; ".join(c.get("prerequisites", [])) or "None"
        if c.get("prerequisite_note"):
            prereq_str += f"  [{c['prerequisite_note']}]"
        meta_rows = [
            ("Credits", str(c["credits"])),
            ("Semester", c.get("semester", "")),
            ("Level", str(c.get("level", ""))),
            ("Prerequisites", prereq_str),
            ("Status", c.get("status", "")),
        ]
        if c.get("icreps_alignment"):
            meta_rows.append(("ICREPs alignment", c["icreps_alignment"]))
        if c.get("embedded_certifications"):
            meta_rows.append(("Embedded certs", ", ".join(c["embedded_certifications"])))

        meta_tbl = doc.add_table(rows=len(meta_rows), cols=2)
        meta_tbl.autofit = False
        _set_col_widths(meta_tbl, [1.8, 4.7])
        _table_default_borders(meta_tbl, color="BFBFBF", sz=4)
        for i, (k, v) in enumerate(meta_rows):
            fill = AMBER_LIGHT_HEX if is_inferred else ALT_ROW_HEX if i % 2 == 0 else "FFFFFF"
            _shade_cell(meta_tbl.rows[i].cells[0], KU_NAVY_HEX)
            _cell_text(meta_tbl.rows[i].cells[0], k, size_pt=SMALL_PT, color=WHITE, bold=True)
            _cell_text(meta_tbl.rows[i].cells[1], v, size_pt=SMALL_PT, color=BODY_GREY)
        _add_gold_left_border(meta_tbl)

        # Description
        _h3(doc, "Description")
        _para(doc, c.get("description", ""), size_pt=BODY_PT, line_spacing=1.25, space_after_pt=6)

        # PLOs
        _h3(doc, "Programme Learning Outcomes (PLOs)")
        for plo_id in sorted(c.get("plos", [])):
            label = plos.get(plo_id, "")
            _bullet(doc, f"{plo_id} — {label}", size_pt=BODY_PT)

        # CLOs
        _h3(doc, "Course Learning Outcomes (CLOs)")
        for clo in c.get("clos", []):
            _bullet(doc, clo, size_pt=BODY_PT)

        # KU GELO
        _h3(doc, "KU GELO Alignment")
        gelo_str = ", ".join(c.get("ku_gelo", []))
        _para(doc, gelo_str, size_pt=SMALL_PT, color=KU_NAVY, space_after_pt=6)

        if is_inferred and c.get("_inference_note"):
            _para(doc, f"Note: {c['_inference_note']}", size_pt=TINY_PT, color=AMBER, italic=True, space_after_pt=8)


# ---------------------------------------------------------------------------
# 5. Quarter × Track × Gate Roadmap Matrix
# ---------------------------------------------------------------------------

def build_roadmap_matrix(doc, data: dict) -> None:
    _h1(doc, "Quarter × Track × Gate Roadmap Matrix")
    _para(doc,
          "The five-quarter implementation plan below spans Q3 2026 through Q3 2027. "
          "Each column identifies deliverables, governance gates, owners, and decision-makers "
          "for both tracks in parallel.",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    phases = data["phases"]
    headers = ["Quarter", "Label", "Minor Deliverable", "Concentration Deliverable", "Gate", "Owner / Decision Maker"]
    tbl = doc.add_table(rows=1 + len(phases), cols=len(headers))
    tbl.autofit = False
    _set_col_widths(tbl, [0.7, 1.0, 1.5, 1.5, 1.2, 1.6])
    _table_default_borders(tbl, color="BFBFBF", sz=4)
    _header_row(tbl, headers)

    for i, phase in enumerate(phases, start=1):
        row = tbl.rows[i]
        vals = [
            phase["q"],
            f"{phase['label']}\n{phase['date_range']}",
            phase["deliverables"]["minor"],
            phase["deliverables"]["concentration"],
            phase["gate"],
            f"{phase['owner']}\nDecision: {phase['decision_maker']}",
        ]
        for j, val in enumerate(vals):
            _cell_text(row.cells[j], val, size_pt=TINY_PT, color=BODY_GREY)
        if i % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, ALT_ROW_HEX)

    _shade_cell(tbl.rows[-1].cells[0], KU_GOLD_HEX)
    _cell_text(tbl.rows[-1].cells[0], phases[-1]["q"], size_pt=TINY_PT, color=KU_NAVY, bold=True)
    _add_gold_left_border(tbl)

    # Key actions per phase
    _h2(doc, "Key Actions by Quarter")
    for phase in phases:
        _h3(doc, f"{phase['q']} — {phase['label']}")
        for action in phase.get("key_actions", []):
            _bullet(doc, action, size_pt=SMALL_PT)


# ---------------------------------------------------------------------------
# 6. Regulator Alignment Matrix
# ---------------------------------------------------------------------------

def build_regulator_matrix(doc, data: dict) -> None:
    _h1(doc, "Regulator Alignment Matrix")
    _para(doc,
          "The following matrix maps UAE regulatory bodies to the four Sports & Fitness Coaching "
          "Concentration courses (SPMD 305–308). Only concentration courses carry direct "
          "professional-registration relevance; the Sports Medicine Minor (SPMD 301–304) "
          "provides foundational knowledge but does not confer fitness-coach credentials.",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    regulators = data["regulators"]
    conc_codes = sorted(["SPMD305", "SPMD306", "SPMD307", "SPMD308"])
    reg_keys = sorted(regulators.keys())

    headers = ["Regulator", "Full Name", "Tier / Level"] + conc_codes
    tbl = doc.add_table(rows=1 + len(reg_keys), cols=len(headers))
    tbl.autofit = False
    _set_col_widths(tbl, [0.75, 1.6, 1.0, 1.0, 1.0, 1.0, 1.0])
    _table_default_borders(tbl, color="BFBFBF", sz=4)
    _header_row(tbl, headers)

    for i, reg_key in enumerate(reg_keys, start=1):
        reg = regulators[reg_key]
        alignment = reg.get("alignment", {})
        row = tbl.rows[i]
        _cell_text(row.cells[0], reg_key, size_pt=SMALL_PT, color=KU_NAVY, bold=True)
        _cell_text(row.cells[1], reg.get("full_name", ""), size_pt=TINY_PT, color=BODY_GREY)
        _cell_text(row.cells[2], reg.get("tier", ""), size_pt=TINY_PT, color=BODY_GREY)
        for j, code in enumerate(conc_codes):
            align_text = alignment.get(code, "")
            cell = row.cells[3 + j]
            if align_text:
                _cell_text(cell, align_text, size_pt=TINY_PT, color=BODY_GREY)
            else:
                _shade_cell(cell, ALT_ROW_HEX)
                _cell_text(cell, "—", size_pt=TINY_PT, color=BODY_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
        if i % 2 == 0:
            for cell in row.cells:
                existing_fill = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                if existing_fill is None:
                    _shade_cell(cell, ALT_ROW_HEX)
    _add_gold_left_border(tbl)


# ---------------------------------------------------------------------------
# 7. PLO × CLO Matrix
# ---------------------------------------------------------------------------

def build_plo_clo_matrix(doc, data: dict) -> None:
    _h1(doc, "PLO × CLO Matrix")
    _para(doc,
          "The following tables show, for each SPMD course, which Course Learning Outcomes (CLOs) "
          "map to each Programme Learning Outcome (PLO). Courses marked PROPOSED are inferred and "
          "require faculty validation.",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    courses = data["courses"]
    plos = data["bscnd_plos"]
    plo_ids = [p["id"] for p in plos]

    for code in sorted(courses.keys()):
        c = courses[code]
        is_inferred = c.get("_inferred", False)
        _h2(doc, f"{code} — {c['title']}" + (" [PROPOSED]" if is_inferred else ""))

        clos = c.get("clos", [])
        course_plos = set(c.get("plos", []))

        # Build CLO × PLO mapping table
        headers = ["CLO"] + plo_ids
        tbl = doc.add_table(rows=1 + len(clos), cols=1 + len(plo_ids))
        tbl.autofit = False
        widths = [3.5] + [0.5] * len(plo_ids)
        _set_col_widths(tbl, widths)
        _table_default_borders(tbl, color="BFBFBF", sz=4)
        _header_row(tbl, headers)

        for i, clo_text in enumerate(clos, start=1):
            row = tbl.rows[i]
            _cell_text(row.cells[0], clo_text, size_pt=TINY_PT, color=BODY_GREY)
            for j, plo_id in enumerate(plo_ids):
                cell = row.cells[1 + j]
                if plo_id in course_plos:
                    _shade_cell(cell, KU_NAVY_HEX)
                    _cell_text(cell, "●", size_pt=TINY_PT, color=WHITE,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    _shade_cell(cell, ALT_ROW_HEX)
                    _cell_text(cell, "", size_pt=TINY_PT, color=BODY_GREY,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_gold_left_border(tbl)
        _para(doc, f"PLOs addressed: {', '.join(sorted(course_plos))}  |  KU GELOs: {', '.join(c.get('ku_gelo', []))}",
              size_pt=TINY_PT, color=KU_NAVY, italic=True, space_after_pt=8)


# ---------------------------------------------------------------------------
# 8. Known Data Issues
# ---------------------------------------------------------------------------

def build_data_issues(doc, data: dict) -> None:
    _h1(doc, "Known Data Issues")
    _para(doc,
          "The following data issues were identified during the pipeline extraction phase. "
          "Each item requires a decision before UGCC submission.",
          size_pt=BODY_PT, line_spacing=1.25, space_after_pt=8)

    for i, issue in enumerate(data.get("known_data_issues", []), start=1):
        _numbered(doc, f"Decision needed: {issue}")


# ---------------------------------------------------------------------------
# 9. Appendix A — Source Corpus Index
# ---------------------------------------------------------------------------

def build_appendix_corpus(doc, data: dict) -> None:
    _h1(doc, "Appendix A — Source Corpus Index")
    _para(doc,
          "The following source documents were used in compiling this roadmap.",
          size_pt=BODY_PT, space_after_pt=8)

    corpus = data.get("corpus_extracts", [])
    if not corpus:
        _para(doc, "(Corpus extract data not available in this data file.)",
              size_pt=SMALL_PT, italic=True, color=BODY_GREY)
        return

    headers = ["ID", "Title / Source", "Type", "Relevance"]
    tbl = doc.add_table(rows=1 + len(corpus), cols=len(headers))
    tbl.autofit = False
    _set_col_widths(tbl, [0.5, 3.5, 1.0, 1.5])
    _table_default_borders(tbl, color="BFBFBF", sz=4)
    _header_row(tbl, headers)

    for i, item in enumerate(corpus, start=1):
        row = tbl.rows[i]
        _cell_text(row.cells[0], str(i), size_pt=TINY_PT, color=KU_NAVY, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row.cells[1], item.get("title", item.get("source", str(item))), size_pt=TINY_PT, color=BODY_GREY)
        _cell_text(row.cells[2], item.get("type", ""), size_pt=TINY_PT, color=BODY_GREY)
        _cell_text(row.cells[3], item.get("relevance", ""), size_pt=TINY_PT, color=BODY_GREY)
        if i % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, ALT_ROW_HEX)
    _add_gold_left_border(tbl)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    with open(DATA_PATH, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_PT)
    style.font.color.rgb = BODY_GREY

    cover_section = doc.sections[0]
    _setup_section(cover_section)
    _add_header_footer(cover_section)

    build_cover(doc, data)

    _add_page_break(doc)
    build_executive_summary(doc, data)

    _add_page_break(doc)
    build_track_structure(doc, data)

    _add_page_break(doc)
    build_course_catalogue(doc, data)

    _add_page_break(doc)
    build_roadmap_matrix(doc, data)

    _add_page_break(doc)
    build_regulator_matrix(doc, data)

    _add_page_break(doc)
    build_plo_clo_matrix(doc, data)

    _add_page_break(doc)
    build_data_issues(doc, data)

    _add_page_break(doc)
    build_appendix_corpus(doc, data)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    size = OUT_PATH.stat().st_size
    print(f"[OK] Wrote: {OUT_PATH}  ({size:,} bytes)")


if __name__ == "__main__":
    main()
