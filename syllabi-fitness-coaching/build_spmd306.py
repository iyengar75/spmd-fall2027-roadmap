"""Build SPMD 306: Group Fitness Instruction syllabus (UGCC-grade).

Mirrors template structure and SPMD301 reference depth.
Sports & Fitness Coaching track — BSCND Sports Medicine Minor.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path("Syllabi_BSCND/SPMD_Syllabi/SPMD306_Syllabus_Consolidated.docx")

doc = Document()

# Set default font (Calibri 11)
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade_cell(cell, hex_fill="D9E2F3"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def heading(text, bold=True, size=11, space_before=6, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def body(text, size=11, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


# ===== Title =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("SPMD 306: Group Fitness Instruction")
run.bold = True
run.font.size = Pt(14)

body("(2 Lecture hours, 2 Laboratory/Studio hours, 3 Credits)")

# Programme placement note
body(
    "Programme: BSc Clinical Nutrition and Dietetics (BSCND) - Sports Medicine Minor, "
    "Sports & Fitness Coaching track. College of Medicine and Health Sciences, Khalifa University."
)

# ===== Prerequisites =====
heading("Prerequisites")
body("SPMD 301: Foundations of Sports Medicine")

# ===== Corequisites =====
heading("Corequisites")
body("None")

# ===== Course Catalog Description =====
heading("Course Catalog Description")
body(
    "This course prepares students to plan, lead, and evaluate safe and effective group "
    "exercise classes across multiple formats including high-intensity interval training (HIIT), "
    "indoor cycling, aqua fitness, mind-body (yoga and Pilates), strength-based group training, "
    "and dance fitness. Students develop competency in choreography construction using 32-count "
    "music phrasing, multimodal cueing (visual, verbal, and tactile), music selection within "
    "appropriate BPM ranges, and copyright-compliant playlist design. Emphasis is placed on "
    "providing modifications and regressions for mixed-ability participants, motivational and "
    "inclusive communication, pre-participation screening, group safety, and emergency response. "
    "Students apply the REPs UAE Code of Ethical Practice throughout, and the course aligns with "
    "ACE-GFI, AFAA Group Fitness, ACSM-GEI, and REPs UAE Level 3 Group Exercise certification "
    "frameworks."
)

# ===== Textbook =====
heading("Textbook")
body(
    "American Council on Exercise. (2024). ACE group fitness instructor handbook (3rd ed.). "
    "American Council on Exercise."
)

# ===== References =====
heading("Reference Materials")
body(
    "C. Kennedy-Armbruster and M. M. Yoke, Methods of Group Exercise Instruction, 4th ed. "
    "Champaign, IL, USA: Human Kinetics, 2020."
)
body(
    "Athletics and Fitness Association of America, AFAA Principles of Group Fitness Instruction. "
    "Sherman Oaks, CA, USA: AFAA, 2022."
)
body(
    "American College of Sports Medicine, ACSM's Resources for the Group Exercise Instructor. "
    "Philadelphia, PA, USA: Wolters Kluwer, 2022."
)
body(
    "ACSM's Health & Fitness Journal. Selected peer-reviewed articles on group exercise "
    "leadership, music phrasing, and class design."
)
body(
    "Register of Exercise Professionals UAE, REPs UAE Code of Ethical Practice and Level 3 Group "
    "Exercise Occupational Standards. Dubai, UAE: REPs UAE, 2023."
)

# ===== Course Structure & Learning Methodology =====
heading("Course Structure & Learning Methodology")
body(
    "Lectures are delivered as 2 x 50-minute sessions per week, complemented by 1 x 100-minute "
    "studio session per week. The course combines interactive lectures, studio-based teach-backs, "
    "peer-coaching critiques, video self-review, and guest demonstrations from certified group "
    "fitness instructors. Students progressively build a class-design portfolio and rehearse "
    "instructor competencies aligned with ACE-GFI, AFAA, ACSM-GEI, and REPs UAE Level 3 standards, "
    "culminating in a full-length practical teaching demonstration assessed against industry "
    "rubrics."
)

# ===== Course Topics =====
heading("Course Topics")
body(
    "The 14-week schedule below sequences instructor competencies from foundational scope-of-practice "
    "and choreography construction through format-specific class design, into participant-centred "
    "modifications and group leadership, and culminates in industry-aligned practical assessment."
)
topics = [
    ("01", "Foundations of group exercise leadership; the REPs UAE scope of practice; ACE-GFI, AFAA, and ACSM-GEI competency frameworks; instructor-participant relationship and class climate"),
    ("02", "Class design and choreography fundamentals; 32-count music phrasing; pyramid, layered, and add-on choreography construction methods; warm-up, conditioning, and cool-down architecture"),
    ("03", "Cueing methods: visual (mirroring, hand signals), verbal (anticipatory, instructional, motivational), and tactile (consent-based touch-cueing) communication strategies and language inclusivity"),
    ("04", "Music selection, BPM ranges by class format (HIIT 140-160, cycling 80-130, aqua 125-150, yoga 60-90, dance 130-150), copyright compliance, and licensed music libraries"),
    ("05", "HIIT class design: work-to-rest ratios (Tabata 20:10, AMRAP, EMOM); intensity progressions, RPE monitoring, contraindications, and safety screening"),
    ("06", "Indoor cycling format: 5-point bike setup, profile design (hill, flat, intervals), resistance and cadence cueing, heart-rate zones, and contraindications"),
    ("07", "Aqua fitness fundamentals: water properties (buoyancy, hydrostatic pressure, drag), shallow and deep-water choreography, pool safety, and aquatic adaptations for arthritis and prenatal participants"),
    ("08", "Mind-body formats: yoga (Sun Salutation A and B sequencing) and Pilates basics (the six principles); breathwork; inclusive language; and trauma-informed instruction"),
    ("09", "Strength-based group training (BodyPump-style): periodisation across muscle groups, tempo cueing, form correction, equipment selection, and safe load progression"),
    ("10", "Dance fitness (Zumba-style): rhythm analysis (4/4 and Latin rhythms), freestyle choreography, transitions, and adapting routines to mixed cultural participants"),
    ("11", "Modifications and regressions for special populations: pregnancy and postpartum, hypertension, type 2 diabetes, joint replacement, older adults, and adolescent participants"),
    ("12", "Group dynamics, motivational climate (mastery vs. ego), self-determination theory, behaviour change strategies, and managing class disruptions"),
    ("13", "Emergency response and risk management: AED operation, choking response, hypoglycaemia, heat illness, syncope, and incident documentation per REPs UAE and KU campus protocols"),
    ("14", "Final teaching demonstrations and industry-aligned practical assessments mapped to REPs UAE Level 3 Group Exercise competency rubrics"),
]
t_topics = doc.add_table(rows=len(topics) + 1, cols=2)
t_topics.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_topics)
hdr = t_topics.rows[0].cells
hdr[0].text = "Week #"
hdr[1].text = "Topics"
for c in hdr:
    shade_cell(c, "D9E2F3")
    for para in c.paragraphs:
        for r in para.runs:
            r.bold = True
for i, (wk, tp) in enumerate(topics, start=1):
    t_topics.rows[i].cells[0].text = wk
    t_topics.rows[i].cells[1].text = tp
t_topics.columns[0].width = Inches(0.9)
t_topics.columns[1].width = Inches(5.5)

# ===== Laboratories / Studio Tutorials =====
heading("Laboratories / Studio Tutorials")
body(
    "Studio tutorials run in parallel with the lecture sequence and culminate in a mock "
    "REPs UAE Level 3 practical assessment in Week 14."
)
labs = [
    ("01", "Studio orientation, equipment safety walkthrough, PAR-Q+ screening practice, and REPs UAE scope-of-practice scenarios applied to common participant cases"),
    ("02", "Phrase-mapping drills: counting 32-count blocks across pop, Latin, electronic, and instrumental tracks; building 4-block warm-up choreography"),
    ("03", "Cueing lab: rehearsed verbal scripts; visual mirroring and hand signals; consent-based tactile cue rotations; recorded peer feedback"),
    ("04", "Playlist construction lab using the BPM-tagged KU music library, licence verification, and class-arc playlist sequencing"),
    ("05", "HIIT teach-back: lead a 10-minute Tabata or AMRAP segment with RPE checks, contraindication screening, and peer critique"),
    ("06", "Indoor cycling lab: 5-point bike fit assessment and lead a 15-minute hill-and-flat profile with cadence and resistance cueing"),
    ("07", "Aqua fitness pool session: shallow-water choreography, buoyancy and resistance modifications, and aquatic warm-up sequencing"),
    ("08", "Mid-semester Studio Teach-Back (assessed): 15-minute full class segment with mandatory video self-review and rubric-based instructor reflection"),
    ("09", "Strength-group lab: track-by-track form coaching, tempo cueing for eccentric and concentric phases, and equipment transition choreography"),
    ("10", "Dance fitness lab: lead a freestyle warm-up and chorus block with rhythm analysis and culturally inclusive song selection"),
    ("11", "Modification clinic: regression and progression options for pregnancy, hypertension, type 2 diabetes, joint replacement, and older adults"),
    ("12", "Class Design Portfolio workshop: cross-format peer review of class plans, choreography notation, and risk-management addenda"),
    ("13", "Emergency simulation: AED operation, choking response, hypoglycaemia management, heat-illness triage, and incident documentation"),
    ("14", "Final Practical Teaching Demonstration assessed against the REPs UAE Level 3 Group Exercise competency rubric"),
]
t_labs = doc.add_table(rows=len(labs) + 1, cols=2)
t_labs.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_labs)
hdr = t_labs.rows[0].cells
hdr[0].text = "Week #"
hdr[1].text = "Studio / Laboratory Activity"
for c in hdr:
    shade_cell(c, "D9E2F3")
    for para in c.paragraphs:
        for r in para.runs:
            r.bold = True
for i, (wk, tp) in enumerate(labs, start=1):
    t_labs.rows[i].cells[0].text = wk
    t_labs.rows[i].cells[1].text = tp
t_labs.columns[0].width = Inches(0.9)
t_labs.columns[1].width = Inches(5.5)

# ===== Weekly CLO Coverage and Required Readings =====
heading("Weekly CLO Coverage and Required Readings")
body(
    "The matrix below links each teaching week to the targeted CLOs and the required reading "
    "from the textbook (ACE GFI Handbook, 3rd ed., 2024) and reference materials, supporting "
    "transparent assessment alignment and outcome traceability."
)
readings = [
    ("01", "1, 5", "ACE GFI Handbook Ch. 1-2; REPs UAE Code of Ethical Practice"),
    ("02", "1, 2", "ACE GFI Handbook Ch. 6; Kennedy-Armbruster & Yoke Ch. 4"),
    ("03", "2, 4", "ACE GFI Handbook Ch. 7; AFAA Principles Module 3"),
    ("04", "1, 2", "ACE GFI Handbook Ch. 8; ACSM Resources Ch. 5 on music selection"),
    ("05", "1, 3", "ACE GFI Handbook Ch. 11 HIIT design; ACSM Resources Ch. 9"),
    ("06", "1, 3", "ACE GFI Handbook Ch. 12 indoor cycling; AFAA Cycling Specialty"),
    ("07", "1, 3", "ACE GFI Handbook Ch. 14 aqua fitness; ACSM Resources Ch. 11"),
    ("08", "2, 4", "ACE GFI Handbook Ch. 13 mind-body; Kennedy-Armbruster & Yoke Ch. 10"),
    ("09", "1, 2", "ACE GFI Handbook Ch. 10 strength group training; AFAA Resistance module"),
    ("10", "2, 4", "ACE GFI Handbook Ch. 15 dance fitness; ACSM Fitness Journal selected"),
    ("11", "3, 5", "ACE GFI Handbook Ch. 17 modifications; ACSM Resources Ch. 13"),
    ("12", "4, 5", "ACE GFI Handbook Ch. 4 behaviour change; Kennedy-Armbruster & Yoke Ch. 6"),
    ("13", "5", "ACE GFI Handbook Ch. 18 emergency response; REPs UAE incident protocol"),
    ("14", "1-5", "Review of all required readings; REPs UAE Level 3 competency rubric"),
]
t_rd = doc.add_table(rows=len(readings) + 1, cols=3)
t_rd.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_rd)
hdr = t_rd.rows[0].cells
for i, h in enumerate(["Week #", "Target CLOs", "Required Reading"]):
    hdr[i].text = h
    shade_cell(hdr[i], "D9E2F3")
    for para in hdr[i].paragraphs:
        for r in para.runs:
            r.bold = True
for i, (wk, clo, rd) in enumerate(readings, start=1):
    t_rd.rows[i].cells[0].text = wk
    t_rd.rows[i].cells[1].text = clo
    t_rd.rows[i].cells[2].text = rd
t_rd.columns[0].width = Inches(0.7)
t_rd.columns[1].width = Inches(1.2)
t_rd.columns[2].width = Inches(4.5)

# ===== Studio Resources =====
heading("Laboratory and Studio Resources")
body(
    "Dedicated group exercise studio, professional sound system, BPM-tagged music library "
    "(KU institutionally licensed), heart-rate telemetry units, full-wall mirrors, exercise mats, "
    "dumbbells, resistance bands, step platforms, indoor cycling bikes, pool access for aqua "
    "sessions, video recording stations for self-review, and the Blackboard Learning Management "
    "System."
)

# ===== CLO -> PLO table =====
heading("Course Learning Outcomes (CLO) and Contributions to Program Level BSc Clinical Nutrition and Dietetics Program Learning Outcomes (PLO)*:")
body("*PLOs can be found in the academic catalog")
body("#Emphasis level: H: High; M: Medium; L: Low; N: Nothing specific")

clos = [
    ("1", "Design safe and effective group fitness classes across multiple formats (HIIT, cycling, aqua, mind-body, strength, and dance) using evidence-based programming principles.", "PLO 2", "H"),
    ("2", "Demonstrate evidence-based cueing, 32-count music phrasing, and choreography construction aligned with ACE-GFI and AFAA instructional standards.", "PLO 4", "H"),
    ("3", "Apply modifications and regressions for mixed-ability participants and special populations to ensure inclusive and safe participation.", "PLO 2", "M"),
    ("4", "Communicate professionally and inclusively to motivate diverse group exercise participants and foster a positive class climate.", "PLO 4", "H"),
    ("5", "Apply the REPs UAE Code of Ethical Practice and execute emergency response protocols within group exercise settings.", "PLO 5", "H"),
]
t_clo = doc.add_table(rows=len(clos) + 1, cols=4)
t_clo.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_clo)
hdr = t_clo.rows[0].cells
for i, h in enumerate(["No.", "CLOs", "PLOs", "Emphasis Level#"]):
    hdr[i].text = h
    shade_cell(hdr[i], "D9E2F3")
    for para in hdr[i].paragraphs:
        for r in para.runs:
            r.bold = True
for i, (n, clo, plo, em) in enumerate(clos, start=1):
    t_clo.rows[i].cells[0].text = n
    t_clo.rows[i].cells[1].text = clo
    t_clo.rows[i].cells[2].text = plo
    t_clo.rows[i].cells[3].text = em
t_clo.columns[0].width = Inches(0.5)
t_clo.columns[1].width = Inches(4.2)
t_clo.columns[2].width = Inches(0.9)
t_clo.columns[3].width = Inches(1.0)

# ===== Assessment =====
heading("Assessment")
body("All course learning outcomes are assessed using the following assessment tools.")

assessments = [
    ("Quiz 1", "7%", "4", "1, 2"),
    ("Quiz 2", "8%", "10", "3, 4"),
    ("Choreography Project", "15%", "6", "1, 2"),
    ("Mid-semester Studio Teach-Back", "20%", "8", "2, 4"),
    ("Class Design Portfolio", "15%", "12", "1, 3, 5"),
    ("Final Practical Teaching Demonstration", "35%", "14", "1-5"),
]
total_pct = sum(int(a[1].replace("%", "")) for a in assessments)
assert total_pct == 100, f"Assessment totals to {total_pct}%, must equal 100%"

t_as = doc.add_table(rows=len(assessments) + 2, cols=4)
t_as.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_as)
hdr = t_as.rows[0].cells
for i, h in enumerate(["Assessment Instruments", "Contribution to Course Grade (%)", "Week #", "CLO(s)"]):
    hdr[i].text = h
    shade_cell(hdr[i], "D9E2F3")
    for para in hdr[i].paragraphs:
        for r in para.runs:
            r.bold = True
for i, row in enumerate(assessments, start=1):
    for j, val in enumerate(row):
        t_as.rows[i].cells[j].text = val
total_row = t_as.rows[len(assessments) + 1].cells
total_row[0].text = "Total"
total_row[1].text = "100%"
total_row[2].text = ""
total_row[3].text = ""
for c in total_row:
    shade_cell(c, "EDEDED")
    for para in c.paragraphs:
        for r in para.runs:
            r.bold = True

# ===== Assessment Rubric Descriptions =====
heading("Assessment Descriptions and Rubric Highlights")
body(
    "Quiz 1 (Week 4) - Closed-book quiz covering REPs UAE scope of practice, 32-count music phrasing, "
    "and cueing taxonomy; multiple choice and short answer."
)
body(
    "Quiz 2 (Week 10) - Closed-book quiz covering modifications for special populations and motivational "
    "communication strategies; case-based short answer."
)
body(
    "Choreography Project (Week 6) - Submit a notated 4-block choreography sequence (warm-up, "
    "conditioning, peak, cool-down) mapped to a licensed 32-count playlist; assessed on musicality, "
    "safety, and cue specification."
)
body(
    "Mid-semester Studio Teach-Back (Week 8) - 15-minute live class segment in a chosen format with "
    "instructor self-reflection using a video recording; assessed on cueing accuracy, music phrasing, "
    "and participant engagement."
)
body(
    "Class Design Portfolio (Week 12) - Cumulative portfolio of three complete class plans across distinct "
    "formats (HIIT, cycling or aqua, and mind-body), each including a participant-screening protocol, "
    "modification matrix, and risk-management addendum."
)
body(
    "Final Practical Teaching Demonstration (Week 14) - Full 30-minute group fitness class delivered to "
    "peers and a faculty panel; assessed against the REPs UAE Level 3 Group Exercise competency rubric "
    "covering class design, cueing, music selection, modifications, safety, and professional conduct."
)

# ===== Grading Scheme =====
heading("Grading Scheme")
body(
    "The official Khalifa University grading system uses letter grades with pluses and minuses. "
    "Passing grades range from A to D; F is failing. Each letter grade is assigned a grade point "
    "on a four-point scale as per the guidelines below:"
)
heading("For undergraduate programs:", bold=True, size=11)

grade_rows = [
    ("Letter Grade", "Grade Point", "Grade Range", "Description"),
    ("A", "4.00", "From 92.5% to 100%", "Excellent"),
    ("A-", "3.70", "From 89.5% to less than 92.5%", "Very Good"),
    ("B+", "3.30", "From 86.5% to less than 89.5%", "Very Good"),
    ("B", "3.00", "From 82.5% to less than 86.5%", "Good"),
    ("B-", "2.70", "From 79.5% to less than 82.5%", "Good"),
    ("C+", "2.30", "From 76.5% to less than 79.5%", "Satisfactory"),
    ("C", "2.00", "From 72.5% to less than 76.5%", "Satisfactory"),
    ("C-", "1.70", "From 69.5% to less than 72.5%", "Less than Satisfactory"),
    ("D", "1.00", "From 59.5% to less than 69.5%", "Poor"),
    ("F", "0.00", "Less than 59.5%", "Fail"),
]
t_g = doc.add_table(rows=len(grade_rows), cols=4)
t_g.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_g)
for i, row in enumerate(grade_rows):
    for j, val in enumerate(row):
        t_g.rows[i].cells[j].text = val
    if i == 0:
        for c in t_g.rows[i].cells:
            shade_cell(c, "D9E2F3")
            for para in c.paragraphs:
                for r in para.runs:
                    r.bold = True

# ===== CLO Assessment Mapping =====
heading("CLO Assessment Mapping")
body(
    "The matrix below confirms that every CLO is assessed by at least two independent instruments, "
    "in line with KU Undergraduate Curriculum Committee assurance-of-learning requirements."
)
clo_map_rows = [
    ("CLO", "Quiz 1", "Quiz 2", "Choreography Project", "Mid-sem Teach-Back", "Class Design Portfolio", "Final Practical Demo"),
    ("CLO 1", "X", "", "X", "", "X", "X"),
    ("CLO 2", "X", "", "X", "X", "", "X"),
    ("CLO 3", "", "X", "", "", "X", "X"),
    ("CLO 4", "", "X", "", "X", "", "X"),
    ("CLO 5", "", "", "", "", "X", "X"),
]
t_map = doc.add_table(rows=len(clo_map_rows), cols=7)
t_map.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_map)
for i, row in enumerate(clo_map_rows):
    for j, val in enumerate(row):
        t_map.rows[i].cells[j].text = val
    if i == 0:
        for c in t_map.rows[i].cells:
            shade_cell(c, "D9E2F3")
            for para in c.paragraphs:
                for r in para.runs:
                    r.bold = True

# ===== PLO Coverage Summary =====
heading("Program Learning Outcome (PLO) Coverage Summary")
body(
    "Aggregated emphasis profile across the five BSCND PLOs, summarising the contribution of SPMD 306 "
    "to programme-level assurance of learning. PLO 4 (Communication and Collaborative Practice) is the "
    "primary contribution; PLO 5 (Ethics and Professional Responsibility) is the secondary contribution."
)
plo_rows = [
    ("PLO", "Title", "Number of CLOs", "Aggregate Emphasis"),
    ("PLO 1", "Problem-Solving with Science", "0", "N"),
    ("PLO 2", "System / Program Design", "2 (CLO 1, 3)", "H/M"),
    ("PLO 3", "Experimentation & Data Analysis", "0", "N"),
    ("PLO 4", "Communication & Collaborative Practice", "2 (CLO 2, 4)", "H (Primary)"),
    ("PLO 5", "Ethics & Professional Responsibility", "1 (CLO 5)", "H (Secondary)"),
]
t_plo = doc.add_table(rows=len(plo_rows), cols=4)
t_plo.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_plo)
for i, row in enumerate(plo_rows):
    for j, val in enumerate(row):
        t_plo.rows[i].cells[j].text = val
    if i == 0:
        for c in t_plo.rows[i].cells:
            shade_cell(c, "D9E2F3")
            for para in c.paragraphs:
                for r in para.runs:
                    r.bold = True

# ===== Industry Certification Alignment =====
heading("Industry Certification Alignment")
body(
    "Course content and assessments are deliberately mapped to four international and regional group "
    "fitness certification frameworks so that graduates leave SPMD 306 with the underpinning knowledge "
    "and demonstrated practical competencies required to sit each examination:"
)
cert_rows = [
    ("Certification", "Awarding Body", "Primary CLO Coverage", "Mapped Course Components"),
    ("ACE-GFI Group Fitness Instructor", "American Council on Exercise", "CLO 1, 2, 4", "Weeks 1-10; Choreography Project; Mid-sem Teach-Back"),
    ("AFAA Group Fitness Instructor", "Athletics and Fitness Association of America", "CLO 1, 2, 3", "Weeks 2-11; Class Design Portfolio"),
    ("ACSM-GEI Group Exercise Instructor", "American College of Sports Medicine", "CLO 1, 3, 5", "Weeks 5-13; Class Design Portfolio; Quiz 2"),
    ("REPs UAE Level 3 Group Exercise", "Register of Exercise Professionals UAE", "CLO 1-5", "All weeks; Final Practical Teaching Demonstration"),
]
t_cert = doc.add_table(rows=len(cert_rows), cols=4)
t_cert.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_borders(t_cert)
for i, row in enumerate(cert_rows):
    for j, val in enumerate(row):
        t_cert.rows[i].cells[j].text = val
    if i == 0:
        for c in t_cert.rows[i].cells:
            shade_cell(c, "D9E2F3")
            for para in c.paragraphs:
                for r in para.runs:
                    r.bold = True

# ===== Instructor Information =====
heading("Instructor Information")
body(
    "Course Coordinator: To be assigned by the Department of Public Health and Health Promotion, "
    "College of Medicine and Health Sciences. Office hours will be communicated via Blackboard during "
    "Week 1 and posted on the studio door. Email correspondence will be answered within two KU working "
    "days. Studio coaches and guest certified instructors will be introduced in Week 1 along with the "
    "academic calendar of teach-back sessions."
)

# ===== Course Policies =====
heading("Course Policies")
body(
    "Attendance: Active studio attendance is essential because practical competencies cannot be "
    "developed asynchronously. Students are expected to attend all studio sessions; absences exceeding "
    "20% of contact hours will trigger the KU attendance withdrawal procedure (AAW). Students must "
    "notify the instructor in advance of any planned absence and are responsible for catching up on "
    "missed cueing and choreography material."
)
body(
    "Studio Dress Code and Safety: Closed-toe athletic footwear (or approved aqua footwear for pool "
    "sessions), modest activewear that permits full range of motion, and tied-back hair are required "
    "in every studio session. Personal water bottles and clean exercise towels must be brought to each "
    "class. Jewellery that may catch on equipment must be removed before participation."
)
body(
    "Late Work and Make-Up Assessments: Late submissions of the Choreography Project and Class Design "
    "Portfolio will be penalised at 10% per calendar day, up to a maximum of three days, after which a "
    "grade of zero will be recorded unless documented medical or compassionate grounds are approved by "
    "the course coordinator. Missed teach-backs and the Final Practical Demonstration may be rescheduled "
    "only with documented medical evidence approved per KU Academic Regulation ACA 1003."
)
body(
    "Pre-Participation Screening: Students with pre-existing medical conditions that may affect "
    "vigorous exercise participation must complete the KU Student Health Services exercise clearance "
    "form before Week 2 studio sessions. All participants will complete a PAR-Q+ at the start of the "
    "semester, which is treated as confidential and used solely to inform individualised modifications."
)
body(
    "Video and Image Use: Studio teach-backs are recorded for the sole purpose of student self-reflection "
    "and faculty rubric-based assessment. Recordings are stored on the KU Blackboard secure media server "
    "and deleted at the end of the semester per KU Data Protection Policy ITL 8000. Students may opt out "
    "of having their image used for any purpose beyond rubric assessment by submitting written notice to "
    "the Course Coordinator before Week 2."
)
body(
    "Student Feedback and Continuous Improvement: A mid-semester anonymous feedback survey is administered "
    "after Week 7 to inform mid-cycle adjustments. The standard KU end-of-semester course evaluation is "
    "administered in Week 14. Aggregated feedback informs the next iteration of the syllabus and is "
    "discussed at the BSCND programme committee."
)
body(
    "Reasonable Accommodations: Students with documented disabilities or learning differences may "
    "request accommodations through the KU Office of Special Needs Support. Approved accommodations "
    "will be implemented in collaboration with the Course Coordinator without compromising the "
    "essential practical competencies required by the REPs UAE Level 3 occupational standard."
)

# ===== Academic Integrity =====
heading("Academic Integrity Statement")
body(
    "“As a Khalifa University student, I will not lie, cheat, steal, or use any unfair means "
    "in academic work and will behave according to university rules and UAE societal norms and "
    "expectations.” Refer to ACA 3500 Academic Integrity Policy for more details."
)

# ===== Copyright and Plagiarism =====
heading("Copyright and Plagiarism")
body(
    "All students must adhere to the university's policies on copyright and plagiarism, including "
    "the ethical use of generative AI tools. Any use of AI-generated content must be properly "
    "attributed and must not violate academic integrity standards."
)

# ===== Closing =====
body("This syllabus must be augmented by a syllabus supplement for students.")

# Save
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))

# Verify
size = OUT.stat().st_size
print(f"Wrote: {OUT}")
print(f"Size:  {size} bytes ({size/1024:.1f} KB)")

# Reload to confirm
chk = Document(str(OUT))
print(f"Paragraphs: {len(chk.paragraphs)}")
print(f"Tables:     {len(chk.tables)}")
print(f"Assessment total: {total_pct}%")
