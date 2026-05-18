#!/usr/bin/env python3
"""Regenerate SPMD 305 / 306 / 308 / 302a syllabi in the exact NUTR401 visual/structural format.

Format reference: Syllabi_BSCND/400-level/NUTR401_Syllabus_Consolidated.docx
PLO model:        5 consolidated BSCND PLOs (per PLO-CLO Mapping Matrix, Jan 27 2026)
Output paths:     Syllabi_BSCND/SPMD_Syllabi/<CODE>_Syllabus_Consolidated.docx
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

OUT_DIR = Path("/mnt/c/Users/TPOTMedXR/Desktop/SPMD_BSCND/Syllabi_BSCND/SPMD_Syllabi")

# -------------------- shared content --------------------
GRADING_ROWS = [
    ("A",  "4.00", "From 92.5% to 100%",                 "Excellent"),
    ("A-", "3.70", "From 89.5% to less than 92.5%",      "Very Good"),
    ("B+", "3.30", "From 86.5% to less than 89.5%",      "Very Good"),
    ("B",  "3.00", "From 82.5% to less than 86.5%",      "Good"),
    ("B-", "2.70", "From 79.5% to less than 82.5%",      "Good"),
    ("C+", "2.30", "From 76.5% to less than 79.5%",      "Satisfactory"),
    ("C",  "2.00", "From 72.5% to less than 76.5%",      "Satisfactory"),
    ("C-", "1.70", "From 69.5% to less than 72.5%",      "Less than Satisfactory"),
    ("D",  "1.00", "From 59.5% to less than 69.5%",      "Poor"),
    ("F",  "0.00", "Less than 59.5%",                    "Fail"),
]

GRADING_DESC = (
    "The official Khalifa University grading system uses letter grades with pluses and minuses. "
    "Passing grades range from A to D; F is failing. Each letter grade is assigned a grade point "
    "on a four-point scale, used to calculate the semester and cumulative grade point average (GPA)."
)

AI_STATEMENT = (
    '"As a Khalifa University student, I will not lie, cheat, steal, or use any unfair means in '
    'academic work and will behave according to university rules and UAE societal norms and '
    'expectations." Refer to the Student Code of Conduct for further details.'
)

COPYRIGHT_STATEMENT = (
    "All students must adhere to the university's policies on copyright and plagiarism, including "
    "the ethical use of generative AI tools. Any use of AI-generated content must be properly "
    "attributed and must comply with the course instructor's permitted-use guidelines. "
    "Unauthorized use of copyrighted materials or undisclosed AI assistance constitutes academic "
    "misconduct and is subject to disciplinary action under the Student Code of Conduct."
)

SUPPLEMENT_LINE = "This syllabus must be augmented by a syllabus supplement for students"

# -------------------- per-course specs --------------------
SYLLABI = [
    # ============================================================
    # SPMD 305 — Exercise Programme Design (3 cr, 3 Lec, 0 Lab)
    # ============================================================
    {
        "code": "SPMD 305",
        "title": "Exercise Programme Design",
        "credit_line": "(3 Lecture hours, 0 Laboratory hours, 3 Credits)",
        "prerequisites": ["SPMD 301", "SPMD 302"],
        "corequisites": [],
        "description": (
            "This course develops the practical and theoretical foundations of evidence-based exercise "
            "programme design across the lifespan. Students conduct comprehensive needs analyses; apply "
            "ACSM risk-stratification and the PAR-Q+; integrate the NASM Optimum Performance Training "
            "(OPT) model and linear, undulating, and block periodization; and manipulate FITT-VP variables "
            "to design cardiorespiratory, resistance, flexibility, and power programmes for healthy adults, "
            "youth, older adults, pregnant clients, and clients with chronic disease. Programme rationale "
            "is communicated in writing and orally, and outcomes are monitored using objective and "
            "subjective metrics. Content aligns with the NASM-CPT, ACE-CPT, NSCA-CSCS, ACSM-CPT, and "
            "REPs UAE Level 3 competency frameworks."
        ),
        "textbook": "[1] G. G. Haff and N. T. Triplett, Eds., Essentials of Strength Training and Conditioning, 5th ed. Champaign, IL: Human Kinetics, 2024.",
        "references": [
            "[2] B. A. Sutton, Ed., NASM Essentials of Personal Fitness Training, 7th ed. Burlington, MA: Jones & Bartlett Learning, 2022.",
            "[3] American College of Sports Medicine, ACSM's Guidelines for Exercise Testing and Prescription, 11th ed. Philadelphia, PA: Wolters Kluwer, 2022.",
            "[4] T. O. Bompa and C. Buzzichelli, Periodization: Theory and Methodology of Training, 6th ed. Champaign, IL: Human Kinetics, 2019.",
        ],
        "methodology": (
            "The course consists of 3 hours of interactive lecture per week. Delivery includes case-based "
            "lectures, programme-design workshops, peer review of student-authored programmes, video "
            "analysis of coaching cues, and guest sessions from REPs UAE certified coaches. Students build "
            "a portfolio of three client programmes across the semester and defend their design choices "
            "in a Week 13–14 oral case defense."
        ),
        "lecture_topics": [
            "Introduction to exercise programme design; scope of practice for fitness coaches",
            "Health screening and risk stratification (PAR-Q+, ACSM stratification, medical clearance)",
            "Movement assessment (NASM OPT, CHEK, FMS) and postural analysis",
            "Cardiorespiratory programme design (FITT-VP, HR zones, ventilatory thresholds)",
            "Resistance training principles (specificity, overload, variation, individualization)",
            "NASM OPT model phases 1–3: stabilization, strength endurance, hypertrophy",
            "NASM OPT model phases 4–5: maximal strength, power development",
            "Periodization models: linear, undulating, block, conjugate",
            "Flexibility and mobility programme design",
            "Programming for body composition and metabolic health",
            "Programming for athletic and sport performance",
            "Special populations I: youth and older adults",
            "Special populations II: pregnancy, chronic disease, return-to-activity",
            "Monitoring, progression, regression, and programme adjustment",
            "Programme defense and course review",
        ],
        "lab_topics": None,  # no lab/studio
        "resources": (
            "Programme-design software (TrainHeroic, BridgeAthletic free tiers), heart-rate calculators, "
            "1RM calculators, ACSM metabolic calculators, Polar HR telemetry, sample programme templates, "
            "and Blackboard Learning Management System."
        ),
        "clos": [
            ("Conduct comprehensive needs analyses including health-history screening, movement assessment, and goal setting to characterize client readiness and risk",
             "1, 2", "H, H"),
            ("Design evidence-based periodized exercise programmes for cardiorespiratory, resistance, flexibility, and power goals using the OPT and linear/undulating/block models",
             "2", "H"),
            ("Apply FITT-VP variables to manipulate programme intensity, volume, frequency, and progression in response to monitoring data",
             "1, 2", "H, H"),
            ("Modify exercise programmes to safely meet the needs of special populations including youth, older adults, pregnant clients, and clients with chronic disease",
             "2, 5", "H, M"),
            ("Communicate programme rationale and progress to clients, coaches, and the interprofessional team in writing and orally",
             "4", "H"),
        ],
        "assessment": [
            ("Quiz 1",                          "7%",  "5",     "1, 3"),
            ("Quiz 2",                          "8%",  "11",    "2, 4"),
            ("Semester Examination",            "20%", "8",     "1, 2, 3"),
            ("Programme Design Project",        "20%", "10–12", "2, 4"),
            ("Client Case Defense",             "20%", "13–14", "2, 4, 5"),
            ("Final Examination",               "25%", "TBA",   "1, 2, 3, 4, 5"),
        ],
    },
    # ============================================================
    # SPMD 306 — Group Fitness Instruction (3 cr, 2 Lec, 2 Studio)
    # ============================================================
    {
        "code": "SPMD 306",
        "title": "Group Fitness Instruction",
        "credit_line": "(2 Lecture hours, 2 Laboratory hours, 3 Credits)",
        "prerequisites": ["SPMD 301"],
        "corequisites": [],
        "description": (
            "This course develops competence in leading safe, effective, and inclusive group fitness "
            "classes across multiple formats including HIIT, indoor cycling, aqua fitness, mind-body, "
            "dance fitness, and strength-based group training. Students learn class design, 32-count music "
            "phrasing, visual/verbal/tactile cueing, modifications and regressions for mixed-ability "
            "participants, motivational communication, group safety screening, and emergency response. "
            "Studio sessions emphasize teach-back practice, peer-coaching critique, and video self-review. "
            "Content aligns with the ACE Group Fitness Instructor (ACE-GFI), AFAA Group Fitness, "
            "ACSM-GEI, and REPs UAE Level 3 Group Exercise competency frameworks, and prepares students "
            "for the practical-assessment component of those credentials."
        ),
        "textbook": "[1] American Council on Exercise, ACE Group Fitness Instructor Handbook, 3rd ed. San Diego, CA: American Council on Exercise, 2024.",
        "references": [
            "[2] C. M. Kennedy-Armbruster and M. M. Yoke, Methods of Group Exercise Instruction, 4th ed. Champaign, IL: Human Kinetics, 2020.",
            "[3] Athletics and Fitness Association of America, AFAA Principles of Group Fitness Instruction, 2022 ed. Sherman Oaks, CA: AFAA, 2022.",
            "[4] American College of Sports Medicine, ACSM's Resources for the Group Exercise Instructor. Philadelphia, PA: Wolters Kluwer, 2022.",
        ],
        "methodology": (
            "The course consists of 2 hours of lecture and 2 hours of studio practice per week. Lectures "
            "cover class-design theory, music and cueing fundamentals, scope of practice, and group "
            "safety. Studio sessions are hands-on: students rotate as instructor, participant, and peer "
            "reviewer; record their teach-backs for self-analysis; and complete a Week-14 mock REPs-UAE "
            "practical assessment graded by a certified instructor."
        ),
        "lecture_topics": [
            "Foundations of group exercise; REPs UAE scope of practice and code of conduct",
            "Class design and choreography fundamentals; 32-count music phrasing",
            "Cueing methods: visual, verbal, and tactile",
            "Music selection, BPM ranges, copyright, and licensed catalogues",
            "HIIT class design and intensity management",
            "Indoor cycling class format and bike fit",
            "Aqua fitness fundamentals and special considerations",
            "Mind-body formats: yoga and Pilates foundations",
            "Strength-based group training (BodyPump-style); equipment safety",
            "Dance fitness (Zumba-style); cultural sensitivity and choreography ethics",
            "Modifications and regressions for mixed-ability participants",
            "Group dynamics, motivation, and behaviour-change techniques",
            "Emergency response, incident reporting, and risk management in group settings",
            "Final teaching demonstrations and instructor portfolio review",
        ],
        "lab_topics": [
            "Studio orientation; participant intake; class-floor safety procedures",
            "Choreography practice; building 32-count combinations",
            "Cueing teach-back; visual/verbal/tactile drill rotations",
            "Music-and-movement matching workshop; BPM-phrase alignment drills",
            "HIIT class teach-back (peer-coached and recorded)",
            "Indoor cycling teach-back; cadence and resistance cues",
            "Aqua fitness practical (pool session at KU Sports Center)",
            "Mid-semester studio teach-back; instructor evaluated",
            "Mind-body teach-back; breath cueing and tactile adjustments",
            "Strength-group teach-back; spotting and equipment management",
            "Dance fitness teach-back; choreography retention practice",
            "Inclusive-class workshop: regressions for older adults, prenatal, beginner",
            "Emergency drill: simulated incident and AED response",
            "Final practical teaching demonstration (mock REPs-UAE assessment)",
        ],
        "resources": (
            "Group exercise studio, sound system, KU-licensed BPM-tagged music library, heart-rate "
            "telemetry, full-length mirrors, mats, dumbbells, kettlebells, step platforms, indoor cycles, "
            "yoga props, pool access at the KU Sports Center, video-review platform, and Blackboard LMS."
        ),
        "clos": [
            ("Design safe, effective, and inclusive group fitness classes across multiple formats including HIIT, cycling, aqua, mind-body, dance, and strength",
             "2", "H"),
            ("Demonstrate evidence-based cueing, music phrasing, and choreography in live teach-back settings",
             "4", "H"),
            ("Apply modifications and regressions to safely include participants of varying ability, age, and clinical status",
             "2, 4", "H, M"),
            ("Communicate professionally and inclusively to motivate diverse groups and build participant adherence",
             "4", "H"),
            ("Apply the REPs UAE Code of Ethical Practice, participant safety screening, and emergency-response protocols in group settings",
             "5", "H"),
        ],
        "assessment": [
            ("Quiz 1",                                  "7%",  "4",  "1, 2"),
            ("Quiz 2",                                  "8%",  "10", "3, 4"),
            ("Choreography Project",                    "15%", "6",  "1, 2"),
            ("Mid-Semester Studio Teach-Back",          "20%", "8",  "2, 4"),
            ("Class Design Portfolio",                  "15%", "12", "1, 3, 5"),
            ("Final Practical Teaching Demonstration",  "35%", "14", "1, 2, 3, 4, 5"),
        ],
    },
    # ============================================================
    # SPMD 308 — Supervised Coaching Practicum (3 cr, 0 Lec, 120 hrs Field)
    # ============================================================
    {
        "code": "SPMD 308",
        "title": "Supervised Coaching Practicum",
        "credit_line": "(0 Lecture hours, 120 Laboratory hours, 3 Credits)",
        "prerequisites": ["SPMD 301", "SPMD 302", "SPMD 305", "SPMD 306"],
        "corequisites": ["Valid CPR/AED + First Aid certification (AHA-BLS or equivalent accepted by REPs UAE)"],
        "description": (
            "This capstone practicum totals 120 contact hours of supervised coaching at KU-approved "
            "partner sites including the KU Sports Center, Fitness First UAE, Gold's Gym Abu Dhabi, "
            "Cleveland Clinic Abu Dhabi Wellness, and corporate-wellness programmes. Students perform "
            "fitness assessments, design and deliver individual and group exercise programmes, coach "
            "clients across the lifespan and ability spectrum, document sessions in compliance with the "
            "Department of Health–Abu Dhabi and REPs UAE professional standards, and complete a weekly "
            "reflective seminar at KU. The course bridges classroom theory (SPMD 301–306) with industry "
            "practice and supplies the supervised practical hours required for the NASM-CPT, ACE-CPT, "
            "NSCA-CSCS, and REPs UAE Level 3 practical-assessment pathways."
        ),
        "textbook": "[1] J. W. Coburn and M. H. Malek, Eds., NSCA's Essentials of Personal Training, 3rd ed. Champaign, IL: Human Kinetics, 2022.",
        "references": [
            "[2] Register of Exercise Professionals UAE, REPs UAE Code of Ethical Practice, current ed. Abu Dhabi: REPs UAE, 2024.",
            "[3] American College of Sports Medicine, ACSM's Guidelines for Exercise Testing and Prescription, 11th ed. Philadelphia, PA: Wolters Kluwer, 2022.",
            "[4] Department of Health Abu Dhabi, Professional Standards for Allied Health Practitioners, current ed. Abu Dhabi: DoH-AD, 2024.",
        ],
        "methodology": (
            "The course consists of 120 supervised on-site coaching hours across 14 weeks (approximately "
            "8–9 hours per week) plus a weekly 50-minute reflective seminar at KU. A REPs-UAE Level 3+ "
            "(or equivalent) site supervisor co-evaluates each student with the KU course director. KU "
            "faculty conduct three formal site visits in Weeks 4, 9, and 13. Students maintain a session "
            "log, build an ePortfolio of three client cases, and present a final case-defense in Week 14."
        ),
        "lecture_topics": [
            "Practicum orientation; site assignment; safeguarding and scope of practice",
            "Client intake, screening (PAR-Q+, ACSM risk stratification), informed consent",
            "Fitness-assessment battery: body composition, cardio, strength, flexibility",
            "Programme design for first clients; KU faculty site visit #1",
            "Coaching cues, exercise demonstration, spotting, and manual assistance",
            "Behaviour-change strategies and motivational interviewing in coaching",
            "Mid-practicum review and supervisor 360° feedback",
            "Group-exercise leadership rotation",
            "Special-population client (older adult OR pre-diabetic OR youth); site visit #2",
            "Outcome re-assessment and programme adjustment",
            "Documentation and data privacy (DoH-AD, PHI handling)",
            "Business of coaching: scope, referrals, insurance, professional registration",
            "Emergency drill and incident-report exercise; site visit #3",
            "Final case presentation and ePortfolio defense",
        ],
        "lab_topics": [
            "On-site coaching: 8 hours (orientation and shadowing)",
            "On-site coaching: 9 hours (first independent screening/intake)",
            "On-site coaching: 9 hours (assessment battery execution)",
            "On-site coaching: 9 hours (initial programme delivery)",
            "On-site coaching: 9 hours (coaching cue practice)",
            "On-site coaching: 9 hours (behaviour-change conversations)",
            "On-site coaching: 8 hours (mid-practicum supervisor review)",
            "On-site coaching: 9 hours (group-class co-instruction)",
            "On-site coaching: 9 hours (special-population client)",
            "On-site coaching: 9 hours (outcome re-assessment)",
            "On-site coaching: 9 hours (documentation and EHR practice)",
            "On-site coaching: 8 hours (referrals and interprofessional handoff)",
            "On-site coaching: 8 hours (emergency drill and incident-report)",
            "On-site coaching: 7 hours (final case wrap-up); 120 hours total",
        ],
        "resources": (
            "KU Sports Center and partner-site facilities (Fitness First UAE, Gold's Gym Abu Dhabi, "
            "Cleveland Clinic Abu Dhabi Wellness, corporate wellness partners); client-management "
            "software (Trainerize, MyFitnessPal Coach); ePortfolio platform (PebblePad or LinkedIn "
            "Learning); standardized KU session-log forms; supervisor-rubric documentation; and "
            "Blackboard LMS."
        ),
        "clos": [
            ("Conduct ethical and safe fitness assessments on real clients in compliance with ACSM and REPs UAE professional standards",
             "1, 5", "H, H"),
            ("Design and adjust individualized exercise programmes for at least three clients of varying ability and goals",
             "2", "H"),
            ("Deliver coaching sessions using evidence-based cueing, spotting, and behaviour-change techniques",
             "2, 4", "H, H"),
            ("Document client sessions in compliance with Department of Health–Abu Dhabi and REPs UAE professional standards",
             "4, 5", "H, H"),
            ("Reflect critically on coaching practice using a structured reflective-practice model",
             "4, 5", "M, H"),
        ],
        "assessment": [
            ("Weekly Session Logs",                                   "15%", "1–14", "1, 2, 4"),
            ("Mid-Practicum Site Supervisor Evaluation",              "15%", "7",    "1, 2, 3, 4"),
            ("Reflective Seminar Participation",                      "10%", "1–14", "5"),
            ("Client Programme Portfolio",                            "15%", "10",   "2"),
            ("Final Site Supervisor Evaluation",                      "20%", "14",   "1, 2, 3, 4"),
            ("Final Case Presentation & ePortfolio Defense",          "25%", "14",   "1, 2, 3, 4, 5"),
        ],
    },
    # ============================================================
    # SPMD 302a — Exercise Physiology Laboratory (1 cr, 0 Lec, 2 Lab)
    # ============================================================
    {
        "code": "SPMD 302a",
        "title": "Exercise Physiology Laboratory",
        "credit_line": "(0 Lecture hours, 2 Laboratory hours, 1 Credit)",
        "prerequisites": ["BIOL 207", "SPMD 301"],
        "corequisites": ["SPMD 302"],
        "description": (
            "This laboratory course is the companion to SPMD 302 (Exercise Physiology) and provides "
            "hands-on practice with the assessment methods that underpin exercise prescription. Students "
            "perform indirect calorimetry and VO₂max testing on metabolic carts, determine lactate "
            "thresholds, measure anaerobic power via the Wingate and vertical-jump protocols, apply "
            "heart-rate and rating-of-perceived-exertion monitoring, conduct body-composition assessment "
            "by skinfolds, bioelectrical impedance, and air-displacement plethysmography, and complete "
            "submaximal aerobic capacity prediction protocols. Each session emphasizes ACSM-compliant "
            "safety, informed consent, and standardized reporting suitable for the practical-skill "
            "components of the ACSM-CPT, NSCA-CSCS, and REPs UAE Level 3 credentials."
        ),
        "textbook": "[1] American College of Sports Medicine, ACSM's Guidelines for Exercise Testing and Prescription, 11th ed. Philadelphia, PA: Wolters Kluwer, 2022.",
        "references": [
            "[2] P. J. Maud and C. Foster, Eds., Physiological Assessment of Human Fitness, 2nd ed. Champaign, IL: Human Kinetics, 2006.",
            "[3] V. H. Heyward and A. L. Gibson, Advanced Fitness Assessment and Exercise Prescription, 7th ed. Champaign, IL: Human Kinetics, 2014.",
            "[4] W. L. Kenney, J. H. Wilmore, and D. L. Costill, Physiology of Sport and Exercise, 8th ed. Champaign, IL: Human Kinetics, 2022.",
        ],
        "methodology": (
            "The course consists of one 2-hour laboratory session per week. Each session includes a "
            "10-minute pre-lab safety briefing, 75 minutes of hands-on assessment with students rotating "
            "through tester, subject, recorder, and spotter roles in small groups, and a 30-minute "
            "data-analysis and lab-notebook block. Lecture content is delivered in the companion course "
            "SPMD 302. The lab is staffed by the lab director plus one graduate teaching assistant per "
            "12 students."
        ),
        "lecture_topics": None,  # pure lab — skip lecture-topics table
        "lab_topics": [
            "Lab safety, informed consent, PAR-Q+, and baseline anthropometrics",
            "Resting metabolic rate via indirect calorimetry",
            "Submaximal aerobic capacity: YMCA cycle ergometer test",
            "Submaximal aerobic capacity: Rockport 1-mile walk test",
            "VO₂max: graded treadmill protocol (Bruce / modified Bruce)",
            "VO₂max: graded cycle ergometer protocol",
            "Lactate threshold determination",
            "Mid-semester practical skills check",
            "Anaerobic power: Wingate 30-second cycle test",
            "Anaerobic power: vertical jump (Sargent and countermovement)",
            "Body composition: skinfolds (7-site Jackson–Pollock)",
            "Body composition: bioelectrical impedance and air-displacement plethysmography",
            "Thermoregulation lab: heat stress and hydration tracking",
            "Final practical examination (assessment chosen by lottery)",
        ],
        "resources": (
            "Metabolic cart (Parvo Medics TrueOne or Cosmed Quark), cycle ergometers (Lode/Monark), "
            "motorized treadmill, portable lactate analyzer (Lactate Pro 2), air-displacement "
            "plethysmograph (BodPod) or DEXA, BIA scale (InBody), Polar HR telemetry chest straps, "
            "Harpenden skinfold calipers, vertical-jump mat (Just Jump or Vertec), environmental "
            "chamber where available, KU ExPhys lab notebook template, statistical software "
            "(SPSS or Jamovi), and Blackboard LMS."
        ),
        "clos": [
            ("Conduct ACSM-compliant exercise-physiology assessments safely and reliably across cardiorespiratory, anaerobic, and body-composition domains",
             "3, 5", "H, M"),
            ("Analyze and interpret physiological data including VO₂, heart rate, lactate, and RPE to characterize the exercise response",
             "1, 3", "H, H"),
            ("Apply bioenergetic and cardiorespiratory principles to predict and explain assessment results across populations",
             "1, 3", "H, H"),
            ("Document laboratory procedures and results in compliance with ACSM and REPs UAE professional standards",
             "4, 5", "M, H"),
            ("Communicate findings effectively in formal lab reports and oral debriefs with peers and supervisors",
             "4", "H"),
        ],
        "assessment": [
            ("Weekly Lab Notebook Entries",        "20%", "1–14", "1, 2, 4"),
            ("Lab Report 1 — VO₂max",              "15%", "6",    "1, 2, 3"),
            ("Lab Report 2 — Body Composition",    "15%", "12",   "1, 2, 3, 5"),
            ("Mid-Semester Practical Skills Check","15%", "8",    "1"),
            ("Lab Safety and Participation",       "5%",  "1–14", "4"),
            ("Final Practical Examination",        "30%", "14",   "1, 2, 3, 4, 5"),
        ],
    },
]

# -------------------- builder --------------------
def add_bold_heading(doc, text, style_name="Normal"):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = True
    return p

def add_para(doc, text, style_name="Normal"):
    p = doc.add_paragraph(text, style=style_name)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Build a Table Grid with header row + data rows."""
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            tbl.rows[ri].cells[ci].text = str(val)
    return tbl

def build_syllabus(spec):
    doc = Document()
    # body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 1-2: Title + credit line
    title = doc.add_paragraph()
    r = title.add_run(f"{spec['code']}: {spec['title']}")
    r.bold = True
    add_para(doc, spec["credit_line"])
    doc.add_paragraph("")

    # Prerequisites
    add_bold_heading(doc, "Prerequisites")
    for prereq in spec["prerequisites"]:
        add_para(doc, prereq)
    doc.add_paragraph("")

    # Corequisites (optional — only render if present)
    if spec["corequisites"]:
        add_bold_heading(doc, "Corequisites")
        for coreq in spec["corequisites"]:
            add_para(doc, coreq)
        doc.add_paragraph("")

    # Course Catalog Description
    add_bold_heading(doc, "Course Catalog Description")
    add_para(doc, spec["description"])
    doc.add_paragraph("")

    # Textbook (IEEE numbered entry, no bullet)
    add_bold_heading(doc, "Textbook")
    add_para(doc, spec["textbook"])
    doc.add_paragraph("")

    # Reference Materials
    add_bold_heading(doc, "Reference Materials")
    for ref in spec["references"]:
        add_para(doc, ref)
    doc.add_paragraph("")

    # Course Structure & Learning Methodology
    add_bold_heading(doc, "Course Structure & Learning Methodology")
    add_para(doc, spec["methodology"])
    doc.add_paragraph("")

    # Course Topics (lecture) — if present
    if spec.get("lecture_topics"):
        add_bold_heading(doc, "Course Topics", style_name="No Spacing")
        doc.add_paragraph("")
        rows = [(f"{i+1:02d}", topic) for i, topic in enumerate(spec["lecture_topics"])]
        add_table(doc, ["Week #", "Topics"], rows)
        doc.add_paragraph("")

    # Laboratories/Tutorials — if present
    if spec.get("lab_topics"):
        add_bold_heading(doc, "Laboratories/Tutorials")
        rows = [(f"{i+1:02d}", topic) for i, topic in enumerate(spec["lab_topics"])]
        add_table(doc, ["Week #", "Topics"], rows)
        doc.add_paragraph("")

    # Resources
    add_bold_heading(doc, "Laboratory and/or Computing/Digital Resources")
    add_para(doc, spec["resources"])
    doc.add_paragraph("")

    # CLO/PLO heading + notes
    clo_heading = (
        "Course Learning Outcomes (CLO) and Contributions Program Level "
        "BSc Clinical Nutrition and Dietetics Program Learning Outcomes (PLO)*:"
    )
    add_bold_heading(doc, clo_heading)
    add_para(doc, "*PLOs can be found in the academic catalog")
    add_para(doc, "#Emphasis level: H: High; M: Medium; L: Low; N: Nothing specific")

    # CLO table
    clo_rows = [(str(i + 1), clo_text, plo, emp) for i, (clo_text, plo, emp) in enumerate(spec["clos"])]
    add_table(doc, ["No.", "CLOs", "PLOs", "Emphasis Level#"], clo_rows)
    doc.add_paragraph("")

    # Assessment
    add_bold_heading(doc, "Assessment")
    add_para(doc, "All course learning outcomes are assessed using the following assessment tools.")
    add_table(doc, ["Assessment Instruments", "Contribution to Course Grade (%)", "Week #", "CLO(s)"],
              spec["assessment"])
    doc.add_paragraph("")

    # Grading Scheme
    add_bold_heading(doc, "Grading Scheme")
    add_para(doc, GRADING_DESC)
    add_para(doc, "For undergraduate programs:", style_name="No Spacing")
    add_table(doc, ["Letter Grade", "Grade Point", "Grade Range", "Description"], GRADING_ROWS)
    doc.add_paragraph("")

    # Academic Integrity
    add_para(doc, "Academic Integrity Statement", style_name="No Spacing").runs[0].bold = True
    add_para(doc, AI_STATEMENT, style_name="No Spacing")
    doc.add_paragraph("")

    # Copyright
    add_para(doc, "Copyright and Plagiarism", style_name="No Spacing").runs[0].bold = True
    add_para(doc, COPYRIGHT_STATEMENT, style_name="No Spacing")
    doc.add_paragraph("")

    # Closing
    add_para(doc, SUPPLEMENT_LINE, style_name="No Spacing")

    return doc

def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary = []
    for spec in SYLLABI:
        # Verify assessment sums to 100%
        total = sum(int(row[1].rstrip("%")) for row in spec["assessment"])
        assert total == 100, f"{spec['code']}: assessment sums to {total}%, not 100%"
        assert len(spec["clos"]) == 5, f"{spec['code']}: expected 5 CLOs, got {len(spec['clos'])}"
        for clo_text, plo, emp in spec["clos"]:
            n_plos = len([p.strip() for p in plo.split(",")])
            n_emps = len([e.strip() for e in emp.split(",")])
            assert n_plos == n_emps, f"{spec['code']}: PLO count != emphasis count ({plo} vs {emp})"

        doc = build_syllabus(spec)
        code_nospace = spec["code"].replace(" ", "")
        out_path = OUT_DIR / f"{code_nospace}_Syllabus_Consolidated.docx"
        doc.save(str(out_path))
        size = out_path.stat().st_size
        summary.append((spec["code"], out_path.name, size, len(doc.paragraphs), len(doc.tables)))

    # Print summary
    print("\n=== REGENERATION SUMMARY (NUTR401 format, 5-PLO consolidated model) ===")
    print(f"{'Code':<12} {'File':<42} {'Size':>8} {'Paras':>6} {'Tables':>7}")
    print("-" * 80)
    for code, fname, size, paras, tables in summary:
        print(f"{code:<12} {fname:<42} {size:>8} {paras:>6} {tables:>7}")
    print(f"\nAll {len(SYLLABI)} syllabi regenerated in {OUT_DIR}")

if __name__ == "__main__":
    main()
