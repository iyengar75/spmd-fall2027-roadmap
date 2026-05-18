"""Build SPMD 305: Exercise Programme Design syllabus (UGCC-grade)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

TEMPLATE = '/mnt/c/Users/TPOTMedXR/Desktop/SPMD_BSCND/Syllabi_BSCND/Course Syllabus Template Updated Dec 2025.docx'
REFERENCE = '/mnt/c/Users/TPOTMedXR/Desktop/SPMD_BSCND/Syllabi_BSCND/SPMD_Syllabi/SPMD301_Syllabus_Consolidated.docx'
OUTPUT = '/mnt/c/Users/TPOTMedXR/Desktop/SPMD_BSCND/Syllabi_BSCND/SPMD_Syllabi/SPMD305_Syllabus_Consolidated.docx'

KU_NAVY = RGBColor(0x00, 0x2D, 0x62)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, size=14, bold=True, color=KU_NAVY, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_para(text, bold=False, size=11, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def style_table(table, header_fill='002D62', header_white=True):
    table.style = 'Table Grid'
    for j, cell in enumerate(table.rows[0].cells):
        shade_cell(cell, header_fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                if header_white:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)


# ===== HEADER =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run('SPMD 305: Exercise Programme Design')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = KU_NAVY
run.font.name = 'Calibri'

credit = doc.add_paragraph()
crun = credit.add_run('(3 Lecture hours, 0 Laboratory hours, 3 Credits)')
crun.font.size = Pt(12)
crun.italic = True
crun.font.name = 'Calibri'

# ===== PREREQUISITES =====
add_heading('Prerequisites')
add_para('SPMD 301: Foundations of Sports Medicine')
add_para('SPMD 302: Exercise Physiology')

# ===== COREQUISITES =====
add_heading('Corequisites')
add_para('None')

# ===== COURSE CATALOG DESCRIPTION =====
add_heading('Course Catalog Description')
add_para(
    'This course develops competencies in evidence-based exercise programme design for healthy adults, '
    'athletes, and special populations. Students conduct comprehensive needs analyses, perform pre-participation '
    'health screening (PAR-Q+, ACSM risk stratification), and execute movement assessments to inform programme '
    'prescription. Core content includes the NASM Optimum Performance Training (OPT) model; linear, undulating, '
    'block, and conjugate periodization frameworks; manipulation of the FITT-VP programme variables (Frequency, '
    'Intensity, Time, Type, Volume, Progression); exercise selection and ordering; intensity and volume management; '
    'recovery and monitoring strategies; and population-specific modifications for youth, older adults, pregnant '
    'clients, and individuals with chronic disease. The course aligns with NASM-CPT, ACE-CPT, ISSA-CPT, ACSM-CPT, '
    'NSCA-CSCS, and REPs UAE Level 3 certification competencies.'
)

# ===== TEXTBOOK =====
add_heading('Textbook')
add_para(
    'G. G. Haff and N. T. Triplett, Eds., Essentials of Strength Training and Conditioning, '
    '5th ed. Champaign, IL, USA: Human Kinetics, 2024.'
)

# ===== REFERENCE MATERIALS =====
add_heading('Reference Materials')
add_para(
    'B. A. Sutton, Ed., NASM Essentials of Personal Fitness Training, '
    '7th ed. Burlington, MA, USA: Jones & Bartlett Learning, 2022.'
)
add_para(
    'American College of Sports Medicine, ACSM’s Guidelines for Exercise Testing and Prescription, '
    '11th ed. Philadelphia, PA, USA: Wolters Kluwer, 2022.'
)
add_para(
    'T. O. Bompa and C. Buzzichelli, Periodization: Theory and Methodology of Training, '
    '6th ed. Champaign, IL, USA: Human Kinetics, 2019.'
)
add_para(
    'Journal of Strength and Conditioning Research. Selected peer-reviewed articles, '
    'Lippincott Williams & Wilkins for the National Strength and Conditioning Association.'
)

# ===== COURSE STRUCTURE =====
add_heading('Course Structure & Learning Methodology')
add_para(
    'Lectures: 3 x 50-minutes per week. The course is delivered through interactive lectures, '
    'programme-design workshops, client case studies, peer review of student-authored exercise programmes, '
    'and guest lectures from REPs UAE certified strength & conditioning coaches and personal trainers. '
    'Workshops emphasize translating physiological principles into prescriptive programme variables, with '
    'iterative feedback on programme drafts. Formative assessment is embedded throughout via low-stakes '
    'design exercises and structured peer critique using rubrics aligned to international certification standards.'
)

# ===== COURSE TOPICS TABLE =====
add_heading('Course Topics')
topics = [
    ('01', 'Introduction to exercise programme design and the needs analysis framework'),
    ('02', 'Pre-participation health screening: PAR-Q+, ACSM risk stratification, medical clearance'),
    ('03', 'Movement assessment: NASM OPT overhead squat assessment and CHEK postural analysis'),
    ('04', 'Cardiorespiratory programming using the FITT-VP framework and HR/VO2 reserve methods'),
    ('05', 'Resistance training principles: specificity, overload, variation, and individualization'),
    ('06', 'NASM OPT model Phases 1–3: stabilization endurance, strength endurance, hypertrophy'),
    ('07', 'NASM OPT model Phases 4–5: maximal strength and power development'),
    ('08', 'Periodization paradigms: linear, undulating (DUP), block, and conjugate models'),
    ('09', 'Flexibility and mobility programming: static, dynamic, PNF, and self-myofascial release'),
    ('10', 'Programming for body composition, fat loss, and lean mass accrual'),
    ('11', 'Programming for athletic performance: speed, agility, plyometrics, and sport-specific transfer'),
    ('12', 'Special populations I: youth athletes, older adults, and long-term athletic development'),
    ('13', 'Special populations II: chronic disease (CVD, T2DM, hypertension), pregnancy and postpartum'),
    ('14', 'Programme monitoring, progression and regression strategies, and load management'),
    ('15', 'Programme review and client case defense; final integration and revision'),
]
t = doc.add_table(rows=len(topics) + 1, cols=2)
t.columns[0].width = Cm(2.5)
t.columns[1].width = Cm(14.0)
hdr = t.rows[0].cells
hdr[0].text = 'Week #'
hdr[1].text = 'Topics'
for i, (wk, topic) in enumerate(topics, start=1):
    t.rows[i].cells[0].text = wk
    t.rows[i].cells[1].text = topic
style_table(t)

# ===== DIGITAL RESOURCES =====
add_heading('Laboratory and/or Computing/Digital Resources')
add_para(
    'Programme design software (TrainHeroic and BridgeAthletic free tiers), heart-rate zone calculators, '
    '1RM estimation calculators (Epley, Brzycki), ACSM metabolic equations and calculators, body composition '
    'and energy-expenditure spreadsheets, video movement-screen tools, and the Blackboard Learning Management System.'
)

# ===== CLO-PLO TABLE =====
add_heading('Course Learning Outcomes (CLO) and Contributions to BSc Clinical Nutrition & Dietetics Program Learning Outcomes (PLO)*:')
add_para('*PLOs can be found in the academic catalog', size=10)
add_para('#Emphasis level: H: High; M: Medium; L: Low; N: Nothing specific', size=10)

clos = [
    ('1',
     'Conduct comprehensive needs analyses by integrating client health-screening data, movement assessments, '
     'training history, and goal-setting to inform evidence-based exercise programme design.',
     'PLO 2', 'H'),
    ('2',
     'Design evidence-based periodized exercise programmes (linear, undulating, block) that systematically '
     'manipulate volume, intensity, and exercise selection across macro-, meso-, and microcycles.',
     'PLO 2', 'H'),
    ('3',
     'Apply the FITT-VP framework and the NASM Optimum Performance Training (OPT) model to prescribe '
     'cardiorespiratory, resistance, flexibility, and power components consistent with international '
     'certification standards (NASM-CPT, ACE-CPT, ACSM-CPT, NSCA-CSCS, REPs UAE Level 3).',
     'PLO 1', 'H'),
    ('4',
     'Modify exercise programmes for special populations — youth, older adults, pregnant clients, and '
     'individuals with chronic disease — by analyzing population-specific contraindications, physiological '
     'considerations, and outcome data.',
     'PLO 3', 'M'),
    ('5',
     'Communicate programme rationale, progression logic, and expected outcomes to clients, coaches, and '
     'interdisciplinary teams through written programme documentation and oral case defenses.',
     'PLO 4', 'M'),
]
ct = doc.add_table(rows=len(clos) + 1, cols=4)
hdr = ct.rows[0].cells
hdr[0].text = 'No.'
hdr[1].text = 'CLOs'
hdr[2].text = 'PLOs'
hdr[3].text = 'Emphasis Level#'
for i, (n, clo, plo, em) in enumerate(clos, start=1):
    ct.rows[i].cells[0].text = n
    ct.rows[i].cells[1].text = clo
    ct.rows[i].cells[2].text = plo
    ct.rows[i].cells[3].text = em
style_table(ct)

# ===== ASSESSMENT TABLE =====
add_heading('Assessment')
add_para('All course learning outcomes are assessed using the following assessment tools.')

assessments = [
    ('Quiz 1', '7%', '5', '1, 3'),
    ('Quiz 2', '8%', '11', '2, 4'),
    ('Midterm Examination', '20%', '8', '1, 2, 3'),
    ('Programme Design Project', '20%', '10–12', '2, 4'),
    ('Client Case Defense', '20%', '13–14', '2, 4, 5'),
    ('Final Comprehensive Examination', '25%', 'Finals', '1–5'),
]
at = doc.add_table(rows=len(assessments) + 2, cols=4)
hdr = at.rows[0].cells
hdr[0].text = 'Assessment Instruments'
hdr[1].text = 'Contribution to Course Grade (%)'
hdr[2].text = 'Week #'
hdr[3].text = 'CLO(s)'
for i, (instr, pct, wk, clo) in enumerate(assessments, start=1):
    at.rows[i].cells[0].text = instr
    at.rows[i].cells[1].text = pct
    at.rows[i].cells[2].text = wk
    at.rows[i].cells[3].text = clo
# Total row
total_row = at.rows[len(assessments) + 1].cells
total_row[0].text = 'Total'
total_row[1].text = '100%'
total_row[2].text = ''
total_row[3].text = ''
for r in total_row[0].paragraphs[0].runs + total_row[1].paragraphs[0].runs:
    r.bold = True
style_table(at)

# ===== GRADING SCHEME =====
add_heading('Grading Scheme')
add_para(
    'The official Khalifa University grading system uses letter grades with pluses and minuses. '
    'Passing grades range from A to D; F is failing. Each letter grade is assigned a grade point on a '
    'four-point scale as per the guidelines below:'
)
add_para('For undergraduate programs:', bold=True)

grades = [
    ('A', '4.00', 'From 92.5% to 100%', 'Excellent'),
    ('A-', '3.70', 'From 89.5% to less than 92.5%', 'Very Good'),
    ('B+', '3.30', 'From 86.5% to less than 89.5%', 'Very Good'),
    ('B', '3.00', 'From 82.5% to less than 86.5%', 'Good'),
    ('B-', '2.70', 'From 79.5% to less than 82.5%', 'Good'),
    ('C+', '2.30', 'From 76.5% to less than 79.5%', 'Satisfactory'),
    ('C', '2.00', 'From 72.5% to less than 76.5%', 'Satisfactory'),
    ('C-', '1.70', 'From 69.5% to less than 72.5%', 'Less than Satisfactory'),
    ('D+', '1.30', 'From 66.5% to less than 69.5%', 'Poor'),
    ('D', '1.00', 'From 60% to less than 66.5%', 'Poor'),
    ('F', '0.00', 'Less than 60%', 'Fail'),
]
gt = doc.add_table(rows=len(grades) + 1, cols=4)
hdr = gt.rows[0].cells
hdr[0].text = 'Letter Grade'
hdr[1].text = 'Grade Point'
hdr[2].text = 'Grade Range'
hdr[3].text = 'Description'
for i, (lg, gp, rng, desc) in enumerate(grades, start=1):
    gt.rows[i].cells[0].text = lg
    gt.rows[i].cells[1].text = gp
    gt.rows[i].cells[2].text = rng
    gt.rows[i].cells[3].text = desc
style_table(gt)

# ===== ACADEMIC INTEGRITY =====
add_heading('Academic Integrity Statement')
add_para(
    '“As a Khalifa University student, I will not lie, cheat, steal, or use any unfair means in academic '
    'work and will behave according to university rules and UAE societal norms and expectations.” '
    'Refer to ACA 3500 Academic Integrity Policy for more details.'
)

# ===== COPYRIGHT & PLAGIARISM =====
add_heading('Copyright and Plagiarism')
add_para(
    'All students must adhere to the university’s policies on copyright and plagiarism, including the '
    'ethical use of generative AI tools. Any use of AI-generated content must be properly attributed and must '
    'not violate academic integrity standards.'
)

# ===== CLOSING =====
add_heading('Programme Design Project — Specification')
add_para(
    'In Weeks 10–12, each student authors a 12-week periodized exercise programme for an assigned client '
    'persona drawn from one of four categories: (a) general-population body-composition client, (b) '
    'recreational endurance athlete, (c) intermediate strength-sport athlete, or (d) older adult with '
    'controlled cardiometabolic disease. Deliverables include: (i) a full needs analysis with PAR-Q+ and '
    'ACSM risk-stratification documentation, (ii) a written movement-assessment report identifying '
    'compensations and proposed corrective strategies, (iii) macrocycle, mesocycle, and microcycle planning '
    'documents with explicit FITT-VP variables, (iv) a session-by-session training log template with load '
    'prescriptions and progression criteria, and (v) a one-page client-facing summary written in plain language. '
    'Programmes are evaluated using a rubric aligned to NASM-CPT, ACSM-CPT, and NSCA-CSCS scope-of-practice '
    'requirements.'
)

add_heading('Client Case Defense — Specification')
add_para(
    'In Weeks 13–14, each student presents and defends the programme designed in the previous block to a '
    'panel comprising the instructor and one or more REPs UAE certified guest practitioners. The 15-minute '
    'oral defense requires the student to justify exercise selection, periodization logic, intensity and '
    'volume prescriptions, monitoring strategy, and any contraindication-driven modifications. Panelists pose '
    'scenario-based questions (e.g., client missed two weeks of training; client reports elevated resting '
    'heart rate; client requests a deload). Performance is scored on technical accuracy, communication '
    'clarity, professional demeanor, and the ability to revise the programme in real time in response to new '
    'information.'
)

add_heading('Alignment with International Certification Standards')
add_para(
    'Course content and assessment competencies are mapped to the published Job Task Analyses and content '
    'outlines of the National Academy of Sports Medicine Certified Personal Trainer (NASM-CPT), the American '
    'Council on Exercise Certified Personal Trainer (ACE-CPT), the International Sports Sciences Association '
    'Certified Personal Trainer (ISSA-CPT), the American College of Sports Medicine Certified Personal Trainer '
    '(ACSM-CPT), the National Strength and Conditioning Association Certified Strength and Conditioning '
    'Specialist (NSCA-CSCS), and Register of Exercise Professionals UAE (REPs UAE) Level 3. Successful '
    'completion of SPMD 305, together with SPMD 301–304, is intended to position graduates of the Sports & '
    'Fitness Coaching track to sit for any of these certification examinations.'
)

add_para('')
closing = add_para('This syllabus must be augmented by a syllabus supplement for students')
for r in closing.runs:
    r.italic = True

doc.save(OUTPUT)

# Verify
import os
size = os.path.getsize(OUTPUT)
d = Document(OUTPUT)
print(f'WRITTEN: {OUTPUT}')
print(f'SIZE: {size} bytes')
print(f'PARAGRAPHS: {len(d.paragraphs)}')
print(f'TABLES: {len(d.tables)}')
# Sanity check assessment sum
total = 7 + 8 + 20 + 20 + 20 + 25
print(f'ASSESSMENT TOTAL: {total}%')
# Verify each CLO mapped to PLO
plos = [row[2] for row in clos]
print(f'CLO->PLO mapping: {plos}')
