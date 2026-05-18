#!/usr/bin/env python3
"""Build SPMD 308: Supervised Coaching Practicum syllabus (UGCC-grade)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/mnt/c/Users/TPOTMedXR/Desktop/SPMD_BSCND/Syllabi_BSCND/SPMD_Syllabi/SPMD308_Syllabus_Consolidated.docx"

doc = Document()

# ---- page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ---- default style ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, size=14, bold=True, color=(0x1F, 0x3A, 0x68)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def shade_cell(cell, hex_color="1F3A68"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def style_header_row(row, hex_color="1F3A68"):
    for cell in row.cells:
        shade_cell(cell, hex_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10.5)


# ============================================================
# 1. HEADER
# ============================================================
title = doc.add_paragraph()
trun = title.add_run("SPMD 308: Supervised Coaching Practicum")
trun.bold = True
trun.font.size = Pt(16)
trun.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

# 2. credits line
hrs = doc.add_paragraph()
hrun = hrs.add_run("(0 Lecture hours, 120 Laboratory/Field hours, 3 Credits)")
hrun.italic = True
hrun.font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# 3. PREREQUISITES
# ============================================================
add_heading("Prerequisites", size=13)
add_para("SPMD 301: Foundations of Sports Medicine")
add_para("SPMD 302: Sports Nutrition & Hydration")
add_para("SPMD 305: Exercise Programme Design")
add_para("SPMD 306: Group Fitness Instruction")

# ============================================================
# 4. COREQUISITES
# ============================================================
add_heading("Corequisites", size=13)
add_para(
    "Valid CPR/AED and First Aid certification (American Heart Association Basic Life Support "
    "[BLS] for Healthcare Providers, or equivalent provider accepted by REPs UAE) must be "
    "current for the entire semester. Acceptable alternatives include the American Red Cross "
    "Adult and Pediatric CPR/AED + First Aid certification, the Emirates Red Crescent Authority "
    "CPR/AED course, and the European Resuscitation Council Basic Life Support certificate. "
    "Students must upload a scanned copy of their current certification card to Blackboard "
    "during Week 1; any student whose certification lapses during the semester will be "
    "withdrawn from on-site coaching duties until re-certification is verified."
)

# ============================================================
# 5. CATALOG DESCRIPTION
# ============================================================
add_heading("Course Catalog Description", size=13)
add_para(
    "Supervised coaching practicum totaling 120 contact hours at Khalifa University-approved "
    "partner sites including the KU Sports Center, commercial fitness facilities (e.g., Fitness "
    "First UAE, Gold's Gym Abu Dhabi), corporate wellness programmes, and youth or community "
    "sport organizations. Students perform pre-participation screening and fitness assessments, "
    "design and deliver individualized and small-group exercise programmes, coach clients across "
    "the lifespan and ability spectrum, and document every session in compliance with REPs UAE "
    "and Department of Health-Abu Dhabi (DoH) professional standards. A weekly 50-minute "
    "reflective seminar at KU integrates field experience with classroom theory from SPMD 301-306 "
    "and prepares students for the practical-assessment components of NASM-CPT, ACE-CPT, "
    "NSCA-CSCS, and REPs UAE Level 3 certifications. This course is the capstone of the Sports "
    "& Fitness Coaching track within the BSc Clinical Nutrition & Dietetics Sports Medicine "
    "Minor and bridges classroom theory with industry practice, positioning graduates for "
    "immediate employment in the Abu Dhabi fitness sector and for sitting external certification "
    "examinations within six months of programme completion."
)

# ============================================================
# 6. TEXTBOOK
# ============================================================
add_heading("Textbook", size=13)
add_para(
    "[1] J. W. Coburn and M. H. Malek, Eds., NSCA's Essentials of Personal Training, "
    "3rd ed. Champaign, IL, USA: Human Kinetics, 2022."
)
add_para(
    "This text is used as the primary reference manual throughout the practicum and aligns "
    "directly with the NSCA-CPT and NSCA-CSCS examination blueprints. Students are expected "
    "to consult chapter-level guidance before each new client engagement, particularly for "
    "screening (Ch. 7), assessment (Ch. 10-13), and programme design (Ch. 14-19)."
)

# ============================================================
# 7. REFERENCE MATERIALS
# ============================================================
add_heading("Reference Materials", size=13)
add_para(
    "[2] Register of Exercise Professionals (REPs) UAE, REPs UAE Code of Ethical Practice, "
    "current ed. Abu Dhabi, UAE: REPs UAE, 2024."
)
add_para(
    "[3] American College of Sports Medicine, ACSM's Guidelines for Exercise Testing and "
    "Prescription, 11th ed. Philadelphia, PA, USA: Wolters Kluwer, 2021."
)
add_para(
    "[4] Department of Health - Abu Dhabi, Professional Standards for Allied Health "
    "Practitioners, current ed. Abu Dhabi, UAE: DoH, 2024."
)
add_para(
    "[5] IDEA Health & Fitness Association, IDEA Code of Ethics for Personal Fitness "
    "Trainers. San Diego, CA, USA: IDEA, 2023."
)
add_para(
    "[6] National Academy of Sports Medicine, NASM Essentials of Personal Fitness Training, "
    "7th ed. Burlington, MA, USA: Jones & Bartlett Learning, 2022."
)
add_para(
    "[7] American Council on Exercise, ACE Personal Trainer Manual: The Ultimate Resource "
    "for Fitness Professionals, 6th ed. San Diego, CA, USA: ACE, 2023."
)

# ============================================================
# 8. COURSE STRUCTURE & LEARNING METHODOLOGY
# ============================================================
add_heading("Course Structure & Learning Methodology", size=13)
add_para(
    "The course is delivered as 120 hours of supervised on-site coaching practice across a "
    "14-week semester (approximately 8-9 hours per week), complemented by a weekly 50-minute "
    "reflective seminar held on the KU campus. Each student is matched with a host site and a "
    "designated site supervisor who holds REPs UAE Level 3 or equivalent (NASM-CPT, ACE-CPT, "
    "NSCA-CSCS, or higher) certification. The site supervisor co-evaluates the student with the "
    "KU course director using standardized rubrics aligned to the NCCA-accredited certification "
    "competencies. The KU course director conducts three structured site visits (Weeks 4, 9, "
    "and 13) to observe live coaching, validate session documentation, audit AED and emergency-"
    "response readiness, and confer with the site supervisor on student progression."
)
add_para(
    "Pedagogically, the practicum is grounded in Kolb's (1984) Experiential Learning Cycle: "
    "concrete experience on the gym floor is followed by reflective observation in the seminar, "
    "abstract conceptualization through readings and supervisor feedback, and active "
    "experimentation in subsequent sessions. The reflective seminar uses Gibbs' (1988) "
    "Reflective Cycle to integrate field experiences with theoretical content from SPMD 301-306, "
    "examine ethical dilemmas, and rehearse the practical-assessment formats used by NASM-CPT, "
    "ACE-CPT, NSCA-CSCS, and REPs UAE Level 3 certifications. Students progress through three "
    "scaffolded phases: observation and shadowing (Weeks 1-3), supervised co-coaching "
    "(Weeks 4-9), and independent coaching with supervisor oversight (Weeks 10-14)."
)

# ============================================================
# 9. COURSE TOPICS - SEMINAR + PRACTICUM SCHEDULE
# ============================================================
add_heading("Course Topics (Weekly Reflective Seminar)", size=13)

topics = [
    ("01", "Orientation, site assignment, safeguarding, and scope-of-practice for fitness professionals in the UAE"),
    ("02", "Client intake, pre-participation screening (PAR-Q+, ACSM risk stratification), and informed consent"),
    ("03", "Fitness assessment battery: body composition, cardiorespiratory, muscular strength and endurance, and flexibility"),
    ("04", "Programme design for first clients: needs analysis, periodization, and FITT-VP application"),
    ("05", "Coaching cues, exercise demonstration technique, spotting, and manual assistance"),
    ("06", "Behavior change strategies: motivational interviewing, SMART goals, and adherence techniques"),
    ("07", "Mid-practicum review and supervisor 360-degree feedback session"),
    ("08", "Group exercise leadership rotation: cueing, formation control, and music tempo selection"),
    ("09", "Coaching a special-population client (older adult OR pre-diabetic OR youth athlete)"),
    ("10", "Outcome re-assessment, programme adjustment, and progression/regression decision-making"),
    ("11", "Documentation, data privacy, and electronic record-keeping under DoH-Abu Dhabi and UAE PHI rules"),
    ("12", "Business of coaching: scope of practice, referrals to allied health, liability, and insurance"),
    ("13", "Emergency drill (cardiac arrest simulation), AED deployment, and incident-report exercise"),
    ("14", "Final case presentation and ePortfolio defense"),
]

t_topics = doc.add_table(rows=1, cols=2)
t_topics.style = 'Light Grid Accent 1'
hdr = t_topics.rows[0].cells
hdr[0].text = "Week #"
hdr[1].text = "Seminar Topic"
style_header_row(t_topics.rows[0])
for wk, top in topics:
    row = t_topics.add_row().cells
    row[0].text = wk
    row[1].text = top

doc.add_paragraph()

# ============================================================
# 10. FIELD PRACTICUM HOURS TABLE
# ============================================================
add_heading("Field Practicum Hours", size=13)
add_para(
    "Each student completes a minimum of 120 supervised coaching contact hours over the "
    "14-week practicum. Hours are logged in the standardized KU Session Log and counter-signed "
    "weekly by the site supervisor. Hours are partitioned across three competency phases: "
    "Observation & Onboarding (Weeks 1-3, 25 hours), Supervised Co-Coaching (Weeks 4-9, 53 "
    "hours), and Independent Coaching with Supervisor Oversight (Weeks 10-14, 42 hours). "
    "Falsified or unverified hours constitute an academic integrity violation under ACA 3500 "
    "and result in automatic course failure.",
    italic=False,
)

# Distribute 120 across 14 weeks: 8 hrs x 6 weeks = 48 + 9 hrs x 8 weeks = 72 -> 120
hours_plan = [
    ("01", 8, "Orientation, site walkthrough, shadowing"),
    ("02", 8, "Observe intakes; co-conduct PAR-Q+ screenings"),
    ("03", 9, "Conduct fitness assessments under direct supervision"),
    ("04", 9, "Begin client programme delivery (KU site visit #1)"),
    ("05", 9, "Lead 1:1 sessions; supervisor co-cueing"),
    ("06", 9, "Apply behavior-change techniques in sessions"),
    ("07", 8, "Mid-practicum supervisor evaluation completed"),
    ("08", 9, "Lead small-group fitness component"),
    ("09", 9, "Coach assigned special-population client (KU site visit #2)"),
    ("10", 9, "Re-assess and progress programmes"),
    ("11", 8, "Complete session documentation audit with supervisor"),
    ("12", 9, "Independent coaching of full client caseload"),
    ("13", 8, "Run emergency drill (KU site visit #3); incident-report drafting"),
    ("14", 8, "Final supervisor evaluation; ePortfolio handover"),
]
assert sum(h for _, h, _ in hours_plan) == 120, "Hours must total 120"

t_hours = doc.add_table(rows=1, cols=3)
t_hours.style = 'Light Grid Accent 1'
hdr = t_hours.rows[0].cells
hdr[0].text = "Week #"
hdr[1].text = "Target On-Site Hours"
hdr[2].text = "Weekly Supervisor Check-In Focus"
style_header_row(t_hours.rows[0])
for wk, hrs_val, note in hours_plan:
    row = t_hours.add_row().cells
    row[0].text = wk
    row[1].text = str(hrs_val)
    row[2].text = note
# Total row
total_row = t_hours.add_row().cells
total_row[0].text = "Total"
total_row[1].text = "120"
total_row[2].text = "Meets REPs UAE Level 3 / NASM / ACE / NSCA practicum-hour requirement"
for c in total_row:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_paragraph()

# ============================================================
# 11. RESOURCES
# ============================================================
add_heading("Laboratory / Computing / Digital Resources", size=13)
add_para(
    "KU Sports Center and approved partner-site facilities, including cardiorespiratory and "
    "resistance-training equipment, free-weight platforms, functional-training rigs, body-"
    "composition measurement tools (bioelectrical impedance analyzers, skinfold calipers, "
    "stadiometers), and AED stations. Digital resources include client-management software "
    "(e.g., Trainerize, MyFitnessPal Coach) for programme delivery and adherence tracking; "
    "the electronic ePortfolio platform (LinkedIn Learning or PebblePad) for case curation "
    "and certification-evidence assembly; the standardized KU Session Log template (paper and "
    "digital) for hour verification; the KU Reflective Practice Journal template; and the "
    "Blackboard Learning Management System for seminar materials, rubric distribution, "
    "supervisor-evaluation submission, and grade reporting. Students must also maintain a "
    "current professional indemnity insurance policy approved by the host site, and a personal "
    "incident-report kit aligned with DoH-Abu Dhabi documentation requirements."
)
add_para(
    "Approved partner-site categories from which the KU course director assigns placements "
    "include: (i) the KU Sports Center; (ii) commercial fitness facilities operating under "
    "REPs UAE-registered management (e.g., Fitness First UAE, Gold's Gym Abu Dhabi); "
    "(iii) corporate wellness programmes within ADNOC, Mubadala, and Abu Dhabi Government "
    "entities; and (iv) youth or community sport organizations supervised by Abu Dhabi Sports "
    "Council (ADSC)-recognized coaches."
)

# ============================================================
# 11.a SITE VISIT PROTOCOL
# ============================================================
add_heading("KU Faculty Site-Visit Protocol", size=13)
add_para(
    "The KU course director conducts three structured site visits per student during the "
    "semester. Each visit lasts approximately 90 minutes and follows a standardized protocol:"
)
add_para(
    "Site Visit 1 (Week 4) — Onboarding Verification. The KU faculty member confirms the "
    "student has completed orientation, signed safeguarding and confidentiality agreements, and "
    "is operating within scope of practice. The site supervisor is interviewed regarding student "
    "punctuality, professionalism, and initial coaching aptitude."
)
add_para(
    "Site Visit 2 (Week 9) — Live Coaching Observation. The KU faculty member observes a full "
    "client session conducted by the student (including warm-up, programme execution, cool-down, "
    "and documentation). Observation is scored against the KU Live Coaching Rubric (10 "
    "competency domains) and shared with the student and supervisor within 72 hours."
)
add_para(
    "Site Visit 3 (Week 13) — Emergency Readiness & Documentation Audit. The KU faculty member "
    "audits the student's session-log file, incident-report knowledge, AED accessibility, and "
    "ePortfolio evidence to confirm readiness for the Week 14 capstone defense."
)

# ============================================================
# 11a. PROFESSIONAL CERTIFICATION MAPPING
# ============================================================
add_heading("Mapping to Professional Certification Practical Requirements", size=13)
add_para(
    "The 120 supervised contact hours and assessment package map onto the practical-"
    "experience prerequisites of four NCCA-accredited and one regional certification accepted "
    "by employers in the UAE fitness sector:"
)

cert_map = [
    ("REPs UAE Level 3 (Personal Trainer)",
     "Minimum 30 hours supervised practical coaching; client case-study portfolio; ethics module.",
     "Met by Weeks 4-9 supervised co-coaching + Client Programme Portfolio + REPs Code module."),
    ("NASM-CPT (National Academy of Sports Medicine)",
     "Recommended 100+ hours supervised practical experience; OPT model application.",
     "Met by full 120 practicum hours; OPT model applied in Client Programme Portfolio."),
    ("ACE-CPT (American Council on Exercise)",
     "Recommended 50+ hours supervised practice; IFT model application; AED-current.",
     "Met by Weeks 4-14 supervised practice; IFT model addressed in seminar Weeks 4-6; AED corequisite."),
    ("NSCA-CSCS (Certified Strength & Conditioning Specialist)",
     "Demonstrated experience designing and supervising strength-conditioning programmes.",
     "Met by Client Programme Portfolio (3 programmes) + Final Case Presentation."),
    ("Abu Dhabi Sports Council Coach Endorsement",
     "Documented supervised coaching with ADSC-recognized coach; incident-response competency.",
     "Met when placement is at an ADSC-recognized site; Week 13 emergency drill validates response."),
]

t_cert = doc.add_table(rows=1, cols=3)
t_cert.style = 'Light Grid Accent 1'
hdr = t_cert.rows[0].cells
hdr[0].text = "Certification"
hdr[1].text = "Practical Requirement"
hdr[2].text = "How SPMD 308 Satisfies It"
style_header_row(t_cert.rows[0])
for cert, req, sat in cert_map:
    row = t_cert.add_row().cells
    row[0].text = cert
    row[1].text = req
    row[2].text = sat

doc.add_paragraph()

# ============================================================
# 12. CLO ↔ PLO MAPPING
# ============================================================
add_heading(
    "Course Learning Outcomes (CLO) and Contributions to Program Level "
    "BSc Clinical Nutrition & Dietetics Program Learning Outcomes (PLO)*:",
    size=12,
)
add_para("*PLOs can be found in the academic catalog", italic=True, size=10)
add_para("#Emphasis level: H: High; M: Medium; L: Low; N: Nothing specific", italic=True, size=10)

clos = [
    ("1",
     "Conduct ethical, safe fitness assessments on real clients in accordance with ACSM "
     "Guidelines and REPs UAE standards.",
     "PLO 5", "H"),
    ("2",
     "Design and progressively adjust individualized exercise programmes for at least three "
     "clients of varying ability across the practicum.",
     "PLO 2", "H"),
    ("3",
     "Deliver coaching sessions using evidence-based cueing, spotting, and behavior-change "
     "techniques appropriate to client goals and population.",
     "PLO 4", "H"),
    ("4",
     "Document client sessions, progress notes, and incident reports in compliance with "
     "DoH-Abu Dhabi and REPs UAE professional and data-privacy standards.",
     "PLO 5", "H"),
    ("5",
     "Reflect critically on coaching practice using a structured reflective-practice model "
     "(Gibbs' Reflective Cycle) and integrate supervisor feedback into subsequent sessions.",
     "PLO 4", "M"),
]

t_clo = doc.add_table(rows=1, cols=4)
t_clo.style = 'Light Grid Accent 1'
hdr = t_clo.rows[0].cells
hdr[0].text = "No."
hdr[1].text = "CLOs"
hdr[2].text = "PLOs"
hdr[3].text = "Emphasis Level#"
style_header_row(t_clo.rows[0])
for no, txt, plo, lvl in clos:
    row = t_clo.add_row().cells
    row[0].text = no
    row[1].text = txt
    row[2].text = plo
    row[3].text = lvl

doc.add_paragraph()

# ============================================================
# 13. ASSESSMENT
# ============================================================
add_heading("Assessment", size=13)
add_para(
    "All course learning outcomes are assessed using the following assessment tools. The "
    "instruments are designed as a competency-based assessment package mirroring the format "
    "of the NASM-CPT, ACE-CPT, NSCA-CSCS, and REPs UAE Level 3 practical examinations. Site "
    "supervisor evaluations contribute 35% of the final grade (15% mid-practicum + 20% final), "
    "ensuring that industry practitioners — not solely academic staff — validate the student's "
    "readiness to coach independently. The Final Case Presentation and ePortfolio Defense "
    "(25%) functions as the capstone assessment, requiring students to defend the design, "
    "delivery, and outcome data of three completed client engagements in front of a panel "
    "comprising the KU course director, the site supervisor, and one external assessor."
)

assessments = [
    ("Weekly Session Logs (signed by site supervisor)", "15%", "1-14", "1, 2, 4"),
    ("Mid-Practicum Site Supervisor Evaluation", "15%", "7", "1, 2, 3, 4"),
    ("Reflective Seminar Participation & Journal", "10%", "1-14", "5"),
    ("Client Programme Portfolio (3 individualized programmes)", "15%", "10", "2"),
    ("Final Site Supervisor Evaluation", "20%", "14", "1, 2, 3, 4"),
    ("Final Case Presentation & ePortfolio Defense", "25%", "14", "1, 2, 3, 4, 5"),
]
total_pct = sum(int(a[1].rstrip("%")) for a in assessments)
assert total_pct == 100, f"Assessments must total 100% (got {total_pct}%)"

t_asm = doc.add_table(rows=1, cols=4)
t_asm.style = 'Light Grid Accent 1'
hdr = t_asm.rows[0].cells
hdr[0].text = "Assessment Instruments"
hdr[1].text = "Contribution to Course Grade (%)"
hdr[2].text = "Week #"
hdr[3].text = "CLO(s)"
style_header_row(t_asm.rows[0])
for inst, pct, wk, clo in assessments:
    row = t_asm.add_row().cells
    row[0].text = inst
    row[1].text = pct
    row[2].text = wk
    row[3].text = clo
total_row = t_asm.add_row().cells
total_row[0].text = "Total"
total_row[1].text = "100%"
total_row[2].text = ""
total_row[3].text = ""
for c in total_row:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_paragraph()

# ============================================================
# 13a. STUDENT RESPONSIBILITIES & PROFESSIONAL CONDUCT
# ============================================================
add_heading("Student Responsibilities & Professional Conduct", size=13)
add_para(
    "Students enrolled in SPMD 308 are guests on host-site premises and represent Khalifa "
    "University at all times. The following responsibilities are non-negotiable conditions of "
    "continued enrollment:"
)
add_para(
    "1. Attendance: 100% of scheduled on-site hours must be completed. Missed sessions must "
    "be made up within two weeks; more than two unexcused absences trigger a programme review."
)
add_para(
    "2. Dress code: Closed-toe athletic footwear and KU-issued or KU-approved practicum polo "
    "with name badge. Personal protective equipment as specified by the host site."
)
add_para(
    "3. Scope of practice: Students will not provide medical diagnosis, dietary prescription, "
    "or rehabilitation services outside the scope of a fitness professional. Any client request "
    "outside scope must be referred to the site supervisor and documented."
)
add_para(
    "4. Confidentiality: All client information is protected under DoH-Abu Dhabi PHI rules and "
    "the host site's privacy policy. Photographs, social-media posts, and case-study materials "
    "require written client consent on the KU Consent Form."
)
add_para(
    "5. Safeguarding: Students working with clients under 18 years of age must complete the "
    "KU Safeguarding Module prior to the first session and must never coach a minor unsupervised."
)
add_para(
    "6. Incident reporting: All injuries, near-misses, equipment failures, and adverse events "
    "must be documented on the host-site Incident Report Form within 24 hours, copied to the "
    "KU course director, and reviewed in the next reflective seminar."
)
add_para(
    "7. Continuing professional development: Students are required to attend at least one "
    "external fitness industry event during the semester (e.g., a REPs UAE-accredited workshop, "
    "an ADSC coaching clinic, or a recognized webinar from NASM/ACE/NSCA) and submit a 300-word "
    "reflection to the ePortfolio."
)

# ============================================================
# 14. GRADING SCHEME
# ============================================================
add_heading("Grading Scheme", size=13)
add_para(
    "The official Khalifa University grading system uses letter grades with pluses and minuses. "
    "Passing grades range from A to D; F is failing. Each letter grade is assigned a grade point "
    "on a four-point scale as per the guidelines below:"
)
add_para("For undergraduate programs:", bold=True)

grades = [
    ("A",  "4.00", "From 92.5% to 100%",                  "Excellent"),
    ("A-", "3.70", "From 89.5% to less than 92.5%",       "Very Good"),
    ("B+", "3.30", "From 86.5% to less than 89.5%",       "Very Good"),
    ("B",  "3.00", "From 82.5% to less than 86.5%",       "Good"),
    ("B-", "2.70", "From 79.5% to less than 82.5%",       "Good"),
    ("C+", "2.30", "From 76.5% to less than 79.5%",       "Satisfactory"),
    ("C",  "2.00", "From 72.5% to less than 76.5%",       "Satisfactory"),
    ("C-", "1.70", "From 69.5% to less than 72.5%",       "Less than Satisfactory"),
    ("D",  "1.00", "From 59.5% to less than 69.5%",       "Poor"),
    ("F",  "0.00", "Less than 59.5%",                     "Fail"),
]
t_g = doc.add_table(rows=1, cols=4)
t_g.style = 'Light Grid Accent 1'
hdr = t_g.rows[0].cells
hdr[0].text = "Letter Grade"
hdr[1].text = "Grade Point"
hdr[2].text = "Grade Range"
hdr[3].text = "Description"
style_header_row(t_g.rows[0])
for lg, gp, rng, desc in grades:
    row = t_g.add_row().cells
    row[0].text = lg
    row[1].text = gp
    row[2].text = rng
    row[3].text = desc

doc.add_paragraph()

# ============================================================
# 14a. SUPERVISOR EVALUATION RUBRIC OVERVIEW
# ============================================================
add_heading("Site Supervisor Evaluation Rubric Overview", size=13)
add_para(
    "Both the mid-practicum (Week 7) and final (Week 14) site supervisor evaluations use the "
    "standardized KU Coaching Practicum Rubric, which scores ten competency domains on a 1-5 "
    "scale: (1) Punctuality and professionalism; (2) Client rapport and communication; "
    "(3) Pre-participation screening accuracy; (4) Fitness assessment technique; (5) Programme "
    "design quality; (6) Exercise demonstration and cueing; (7) Spotting and safety supervision; "
    "(8) Behavior change facilitation; (9) Documentation completeness and accuracy; and "
    "(10) Adherence to scope of practice and professional ethics. The supervisor adds a "
    "narrative comment for each domain and a global recommendation regarding the student's "
    "readiness to coach independently in the UAE fitness sector."
)

# ============================================================
# 15. ACADEMIC INTEGRITY
# ============================================================
add_heading("Academic Integrity Statement", size=13)
add_para(
    "“As a Khalifa University student, I will not lie, cheat, steal, or use any unfair "
    "means in academic work and will behave according to university rules and UAE societal "
    "norms and expectations.” Refer to ACA 3500 Academic Integrity Policy for more details."
)

# ============================================================
# 16. COPYRIGHT
# ============================================================
add_heading("Copyright and Plagiarism", size=13)
add_para(
    "All students must adhere to the university’s policies on copyright and plagiarism, "
    "including the ethical use of generative AI tools. Any use of AI-generated content must be "
    "properly attributed and must not violate academic integrity standards. In SPMD 308, AI "
    "tools may not be used to fabricate session data, generate fictitious client interactions, "
    "or compose reflective journal entries describing experiences the student did not actually "
    "have. AI-assisted preparation of programme templates is permitted provided the student "
    "discloses the tool, the prompt, and the manual modifications applied."
)

# ============================================================
# 17. SUPPLEMENT CLOSING
# ============================================================
doc.add_paragraph()
add_para(
    "This syllabus must be augmented by a syllabus supplement for students.",
    italic=True,
)

doc.save(OUT)

# verification
import os
size = os.path.getsize(OUT)
doc2 = Document(OUT)
print(f"OK  file={OUT}")
print(f"    size={size} bytes ({size/1024:.1f} KB)")
print(f"    paragraphs={len(doc2.paragraphs)}")
print(f"    tables={len(doc2.tables)}")
print(f"    assessment_total={total_pct}%")
print(f"    practicum_hours={sum(h for _,h,_ in hours_plan)}")
